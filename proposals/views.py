import io
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from datetime import datetime
from decimal import Decimal
from pathlib import Path

from django.conf import settings
from django.contrib.auth.decorators import login_required
from django.db.models import Max
from django.http import FileResponse, Http404, JsonResponse
from django.contrib import messages
from django.shortcuts import redirect, render
from django.urls import reverse
from django.views.decorators.http import require_http_methods
import markdown
import mammoth
try:
    import pypandoc_binary as pypandoc
except ImportError:
    import pypandoc
from docxtpl import DocxTemplate

from .forms import ComplexProposalForm, ContractForm, ProposalForm, RequisitesParseForm, SROK_CHOICES, TariffForm
from .models import (
    ContractRecord,
    Counterparty,
    KanbanBoardOrder,
    KanbanCardField,
    KanbanCardPosition,
    KanbanColumnCustom,
    KanbanColumnTitleOverride,
    Region,
    RegionServicePrice,
    Service,
    TKPRecord,
)
from .requisites_parser import FIELD_ORDER, parse_requisites_file

COMPLEX_TEMPLATE_NAME = 'Шаблон 9 Комплексное ТКП.docx'
UNIT_DISPLAY = {'m2': 'м²', 'piece': 'шт'}

# Подпапка с шаблонами договоров (внутри TEMPLATES_DOCX_DIR)
CONTRACT_TEMPLATES_SUBDIR = 'contracts_templates'

# Соответствие услуги ТКП → файл шаблона договора (в templates_docx/contracts_templates/)
# Услуга — Шаблон ТКП — Шаблон договора (см. README в contracts_templates)
SERVICE_TO_CONTRACT_TEMPLATE = {
    'ДП': '01_Договор_ДП.docx',
    'ДКП': '02_Договор_ДК.docx',
    'Навигация': '03_Договор_Навигация.docx',
    'Контент': '04_Договор_Контент_система.docx',
    'Навигация_стенды': '05_Договор_Контент_Навигация.docx',
    'Фасад': '06_Договор_ДП_Фасад.docx',
    'ДК Фасад': '07_Договор_ДК_Фасад.docx',
    'Благоустройство': '08_Договор_ДПФ_Благоустройство.docx',
}

# Плейсхолдер в шаблонах договоров 08 и 05 — сюда вставляется таблица спецификации (как в комплексном ТКП)
CONTRACT_SPEC_TABLE_PLACEHOLDER = '___CONTRACT_SPEC_TABLE___'


def get_contract_template_for_complex_tkp(rows):
    """
    Определяет шаблон договора для комплексного ТКП по составу строк.
    - 08: если есть обе услуги Фасад и Благоустройство.
    - 05: если есть не менее двух из трёх: Навигация, Контент, Навигация_стенды.
    Возвращает имя файла шаблона или None.
    """
    if not rows:
        return None
    service_names = {r.get('service_name', '').strip() for r in rows}
    has_facade = 'Фасад' in service_names
    has_blag = 'Благоустройство' in service_names
    if has_facade and has_blag:
        return '08_Договор_ДПФ_Благоустройство.docx'
    nav_content_count = sum(1 for s in ('Навигация', 'Контент', 'Навигация_стенды') if s in service_names)
    if nav_content_count >= 2:
        return '05_Договор_Контент_Навигация.docx'
    return None


def _complex_rows_json_to_ctx(rows_json):
    """Преобразует rows_json (из TKPRecord) в rows_ctx для _build_complex_table_document. Возвращает (rows_ctx, total_sum_formatted)."""
    if not rows_json:
        return [], '0'
    rows_ctx = []
    total = Decimal(0)
    for r in rows_json:
        total += Decimal(str(r.get('total') or 0))
        rows_ctx.append({
            'service_name': r.get('service_name', ''),
            'comment': r.get('comment', ''),
            'srok': r.get('srok', ''),
            'unit_display': UNIT_DISPLAY.get(r.get('unit', ''), r.get('unit', '')),
            'quantity': str(r.get('quantity', '')),
            'price_per_unit': _format_price(Decimal(str(r.get('price_per_unit') or 0))),
            'total_formatted': _format_price(Decimal(str(r.get('total') or 0))),
        })
    return rows_ctx, _format_price(total)

# Текст условий оплаты по умолчанию для договора
DEFAULT_PAYMENT_TERMS = """2.2.1. В течение 5 (пяти) рабочих дней на основании выставленного Исполнителем счета Заказчик выплачивает Исполнителю аванс в размере 30% от цены Договора;
2.2.2. В течение 5 (пяти) рабочих дней после приемки/утверждения результатов работ Заказчик выплачивает Исполнителю оставшиеся 70% цены Договора."""

# Отображаемые названия услуг в форме «Комплексное ТКП» (колонка «Компонент услуги»)
COMPLEX_SERVICE_DISPLAY_NAMES = {
    'ДП': 'Дизайн-проект',
    'ДКП': 'Дизайн-концепция',
    'Навигация': 'Навигация',
    'Контент': 'Контент-система',
    'Навигация_стенды': 'Контент и навигация',
    'Фасад': 'Дизайн-проект Фасада',
    'ДК Фасад': 'Дизайн-концепция Фасада',
    'Благоустройство': 'Благоустройство',
}

# Символ в тексте комментария, по которому делается перенос строки в ТКП и на странице подтверждения
COMPLEX_COMMENT_LINE_BREAK_MARKER = '|'

# Комментарий по умолчанию (fallback, если нет в БД и нет в data/complex_service_comments.json)
COMPLEX_SERVICE_DEFAULT_COMMENTS = {
    'ДП': 'Разработка дизайн-проекта в соответствии с техническим заданием.',
    'ДКП': 'Разработка дизайн-концепции.',
    'Навигация': 'Разработка навигационной системы.',
    'Контент': 'Разработка контент-системы.',
    'Навигация_стенды': 'Контент и навигация для стендов.',
    'Фасад': 'Дизайн-проект фасада.',
    'ДК Фасад': 'Дизайн-концепция фасада.',
    'Благоустройство': 'Проектирование благоустройства территории.',
}

COMPLEX_SERVICE_COMMENTS_FILE = 'data/complex_service_comments.json'


def _load_complex_service_comments_file():
    """Загружает комментарии по умолчанию из файла в папке проекта (для деплоя без ручного ввода)."""
    path = Path(settings.BASE_DIR) / COMPLEX_SERVICE_COMMENTS_FILE
    if not path.exists():
        return {}
    try:
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def _save_complex_service_comments_file(comments_by_name):
    """Сохраняет комментарии в data/complex_service_comments.json (чтобы при деплое не вводить заново)."""
    path = Path(settings.BASE_DIR) / COMPLEX_SERVICE_COMMENTS_FILE
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(comments_by_name, f, ensure_ascii=False, indent=2)


def _get_libreoffice_path():
    """Путь к LibreOffice: на Windows ищем soffice.exe в стандартных каталогах."""
    if sys.platform == 'win32':
        for base in (os.environ.get('ProgramFiles', 'C:\\Program Files'),
                     os.environ.get('ProgramFiles(x86)', 'C:\\Program Files (x86)')):
            exe = Path(base) / 'LibreOffice' / 'program' / 'soffice.exe'
            if exe.exists():
                return str(exe)
    return 'libreoffice'


def _convert_docx_to_pdf(docx_path, out_dir):
    """Конвертация docx в PDF: LibreOffice (Linux/Windows) или docx2pdf (Windows + MS Word)."""
    docx_path = Path(docx_path)
    out_dir = Path(out_dir)
    cmd = _get_libreoffice_path()
    if sys.platform == 'win32' and cmd == 'libreoffice':
        try:
            import pythoncom
            pythoncom.CoInitialize()
        except ImportError:
            pass
        try:
            from docx2pdf import convert as docx2pdf_convert
            docx2pdf_convert(str(docx_path), str(out_dir / 'tkp.pdf'))
            return
        except ImportError:
            raise FileNotFoundError(
                'На Windows нужен LibreOffice или пакет docx2pdf (pip install docx2pdf, требуется MS Word). '
                'Установите LibreOffice или выполните: pip install docx2pdf'
            )
    subprocess.run(
        [cmd, '--headless', '--convert-to', 'pdf', '--outdir', str(out_dir), str(docx_path)],
        check=True,
    )


CONTRACT_FONT_NAME = 'Times New Roman'


def _set_contract_doc_font_times_new_roman(doc):
    """Устанавливает шрифт Times New Roman для всего документа договора (параграфы и ячейки таблиц)."""
    docx = doc.docx
    for paragraph in docx.paragraphs:
        for run in paragraph.runs:
            run.font.name = CONTRACT_FONT_NAME
    for table in docx.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = CONTRACT_FONT_NAME


def _dolznost_from_customer_in_person(customer_in_person):
    """
    Возвращает должность (именительный падеж) по значению поля «Заказчик в лице».

    customer_in_person:
      - «Директора»             -> «Директор»
      - «Генерального директора» -> «Генеральный директор»
      - «И.О. Директора»        -> «И.о. Директора»
      - «Подписанта»            -> «Подписант»
    """
    value = (customer_in_person or '').strip()
    mapping = {
        'Директора': 'Директор',
        'Генерального директора': 'Генеральный директор',
        'И.О. Директора': 'И.о. Директора',
        'Подписанта': 'Подписант',
    }
    return mapping.get(value, '')


def _sanitize_filename(name):
    """Очистка строки для использования в имени файла."""
    if not name or not name.strip():
        return 'client'
    s = name.strip()
    s = re.sub(r'[\s]+', '_', s)
    s = re.sub(r'[\\/:*?"<>|]', '', s)
    return s[:100] if s else 'client'


def _get_next_seq_for_date(date_obj):
    """Порядковый номер ТКП на указанную дату."""
    count = TKPRecord.objects.filter(date=date_obj).count()
    return count + 1


def _get_next_contract_seq_for_date(date_obj):
    """Порядковый номер договора за указанную дату (дата_порядковый). Учитываются все записи за дату (в т.ч. черновики), чтобы номер был уникальным."""
    count = ContractRecord.objects.filter(date=date_obj).count()
    return count + 1


def _get_next_contract_draft_seq_for_date(date_obj):
    """Порядковый номер черновика договора за дату."""
    return ContractRecord.objects.filter(date=date_obj, status=ContractRecord.STATUS_DRAFT).count() + 1


def _format_price(value):
    """Форматирование цены с разделителями тысяч (1 000 000)."""
    if value is None:
        return ''
    s = str(int(value))
    n = len(s)
    if n <= 3:
        return s
    r = n % 3 or 3
    result = s[:r]
    for i in range(r, n, 3):
        result += ' ' + s[i:i + 3]
    return result


def _build_proposal_data_from_form_cleaned(data):
    """Собрать proposal_data из очищенных данных формы (полная валидация уже пройдена)."""
    service = data['service']
    s = data.get('s') or 0
    if data.get('is_internal'):
        price_value = data.get('internal_price') or 0
        region_name = ''
    else:
        region = data['region']
        try:
            rsp = RegionServicePrice.objects.get(region=region, service=service)
            price_value = float(rsp.unit_price) * float(s)
        except RegionServicePrice.DoesNotExist:
            return None, f'Не найдена цена для региона "{region.name}" и услуги "{service.name}".'
        region_name = region.name
    client_value = (
        (data.get('internal_client') or '').strip()
        if data.get('is_internal')
        else (data.get('client') or '')
    )
    s_val = data.get('s')
    if s_val is not None and s_val != '':
        s_str = str(s_val)
    elif data.get('is_internal'):
        s_str = ''
    else:
        s_str = str(s)
    return {
        'date': data['date'].strftime('%Y-%m-%d'),
        'service_id': service.pk,
        'service_name': service.name,
        'city': region_name,
        'price': str(price_value),
        'client': client_value,
        'room': data.get('room') or '',
        'srok': data.get('srok') or '',
        'text': data.get('text') or '',
        's': s_str,
    }, None


@login_required
@require_http_methods(['GET'])
def start_view(request):
    """Дашборд: сводки по ТКП и договорам, последние действия."""
    from django.db.models import Sum
    tkp_draft_count = TKPRecord.objects.filter(status=TKPRecord.STATUS_DRAFT).count()
    tkp_final_count = TKPRecord.objects.filter(status=TKPRecord.STATUS_FINAL).count()
    contract_draft_count = ContractRecord.objects.filter(status=ContractRecord.STATUS_DRAFT).count()
    contract_total_count = ContractRecord.objects.count()
    sum_tkp = TKPRecord.objects.filter(status=TKPRecord.STATUS_FINAL).aggregate(s=Sum('sum_total'))['s'] or 0
    sum_contracts = ContractRecord.objects.filter(status=ContractRecord.STATUS_FINAL).aggregate(s=Sum('sum_total'))['s'] or 0
    recent_tkp = TKPRecord.objects.select_related('created_by').order_by('-created_at')[:10]
    recent_contracts = ContractRecord.objects.select_related('created_by', 'tkp').order_by('-created_at')[:10]
    contract_by_tkp = {
        c.tkp_id: c
        for c in ContractRecord.objects.filter(
            status=ContractRecord.STATUS_FINAL, tkp__isnull=False
        ).select_related('tkp')
    }
    contract_draft_by_tkp = {
        c.tkp_id: c
        for c in ContractRecord.objects.filter(
            status=ContractRecord.STATUS_DRAFT, tkp__isnull=False
        ).select_related('tkp')
    }
    contract_template_by_tkp = {}
    for r in recent_tkp:
        if r.service == 'Комплексное ТКП' and r.rows_json:
            tpl = get_contract_template_for_complex_tkp(r.rows_json)
            contract_template_by_tkp[r.pk] = tpl if tpl else True
    context = {
        'tkp_draft_count': tkp_draft_count,
        'tkp_final_count': tkp_final_count,
        'contract_draft_count': contract_draft_count,
        'contract_total_count': contract_total_count,
        'sum_tkp': sum_tkp,
        'sum_contracts': sum_contracts,
        'recent_tkp': recent_tkp,
        'recent_contracts': recent_contracts,
        'contract_template_by_service': SERVICE_TO_CONTRACT_TEMPLATE,
        'contract_template_by_tkp': contract_template_by_tkp,
        'contract_by_tkp': contract_by_tkp,
        'contract_draft_by_tkp': contract_draft_by_tkp,
    }
    return render(request, 'proposals/start.html', context)


@login_required
@require_http_methods(['GET'])
def instruction_view(request):
    """Страница инструкции пользователя: содержимое docs/ИНСТРУКЦИЯ_ПОЛЬЗОВАТЕЛЯ.md в Markdown."""
    path = settings.BASE_DIR / 'docs' / 'ИНСТРУКЦИЯ_ПОЛЬЗОВАТЕЛЯ.md'
    if path.exists():
        content = path.read_text(encoding='utf-8')
        instruction_html = markdown.markdown(content, extensions=['extra'])
    else:
        instruction_html = '<p>Файл инструкции не найден.</p>'
    return render(request, 'proposals/instruction.html', {'instruction_html': instruction_html})


@login_required
@require_http_methods(['GET', 'POST'])
def form_view(request):
    """Шаг 1: форма ввода параметров ТКП. GET ?draft_id= — возобновление черновика (простое ТКП)."""
    if request.method == 'POST':
        form = ProposalForm(request.POST)
        save_draft = request.POST.get('save_draft')

        if save_draft:
            # Черновик: минимум — дата, услуга, регион (или внутренний клиент + цена)
            draft_errors = []
            date_val = form.data.get('date')
            if not date_val:
                draft_errors.append('Укажите дату.')
            else:
                try:
                    date_obj = datetime.strptime(date_val, '%Y-%m-%d').date()
                except ValueError:
                    draft_errors.append('Некорректная дата.')
                    date_obj = None
            service_id = form.data.get('service')
            service = None
            if service_id:
                try:
                    service = Service.objects.get(pk=service_id)
                except (Service.DoesNotExist, ValueError):
                    pass
            if not service:
                draft_errors.append('Выберите услугу.')
            is_internal = form.data.get('is_internal') == 'on'
            region = None
            if not is_internal:
                region_id = form.data.get('region')
                if region_id:
                    try:
                        region = Region.objects.get(pk=region_id)
                    except (Region.DoesNotExist, ValueError):
                        pass
                if not region:
                    draft_errors.append('Выберите регион.')
            else:
                internal_client = (form.data.get('internal_client') or '').strip()
                if not internal_client:
                    draft_errors.append('Выберите внутреннего клиента.')

            if not draft_errors and date_obj and (service and (region or is_internal)):
                # Собираем данные из POST (любые заполненные поля)
                client_value = ''
                price_value = Decimal(0)
                region_name = ''
                if is_internal:
                    client_value = (form.data.get('internal_client') or '').strip()
                    try:
                        price_value = Decimal(str(form.data.get('internal_price') or 0))
                    except Exception:
                        price_value = Decimal(0)
                else:
                    client_value = (form.data.get('client') or '').strip()
                    region_name = region.name if region else ''
                    s_val = form.data.get('s')
                    try:
                        s_float = float(s_val or 0)
                    except (TypeError, ValueError):
                        s_float = 0
                    if region and service and s_float >= 0:
                        try:
                            rsp = RegionServicePrice.objects.get(region=region, service=service)
                            price_value = rsp.unit_price * Decimal(str(s_float))
                        except RegionServicePrice.DoesNotExist:
                            price_value = Decimal(0)
                s_draft = form.data.get('s')
                proposal_data = {
                    'date': date_obj.strftime('%Y-%m-%d'),
                    'service_id': service.pk,
                    'service_name': service.name,
                    'city': region_name,
                    'price': str(price_value),
                    'client': client_value,
                    'room': (form.data.get('room') or '').strip(),
                    'srok': (form.data.get('srok') or '').strip(),
                    'text': (form.data.get('text') or '').strip(),
                    's': str(s_draft) if s_draft is not None and s_draft != '' else '',
                }
                _save_tkp_record(proposal_data, status=TKPRecord.STATUS_DRAFT, user=request.user)
                messages.success(request, 'Черновик ТКП сохранён в перечень.')
                return redirect('proposals:table')
            if draft_errors:
                for err in draft_errors:
                    messages.error(request, err)
        elif form.is_valid():
            data = form.cleaned_data
            proposal_data, err = _build_proposal_data_from_form_cleaned(data)
            if err:
                messages.error(request, err)
                return render(request, 'proposals/form.html', {
                    'form': form,
                    'service_units_json': json.dumps({str(s.pk): s.unit_type for s in Service.objects.all()}),
                })
            request.session['proposal_data'] = proposal_data
            return redirect('proposals:confirm')
    else:
        draft_id = request.GET.get('draft_id')
        initial = None
        if draft_id:
            try:
                draft = TKPRecord.objects.get(pk=draft_id, status=TKPRecord.STATUS_DRAFT)
                if draft.service != 'Комплексное ТКП':
                    service = Service.objects.filter(name=draft.service).first()
                    initial = {
                        'date': draft.date,
                        'client': draft.client or '',
                        'room': draft.room or '',
                        's': draft.s or '',
                        'text': draft.text or '',
                    }
                    if service:
                        initial['service'] = service.pk
                    request.session['tkp_draft_id'] = int(draft_id)
            except (TKPRecord.DoesNotExist, ValueError, TypeError):
                request.session.pop('tkp_draft_id', None)
        else:
            request.session.pop('tkp_draft_id', None)
        form = ProposalForm(initial=initial) if initial else ProposalForm()

    service_units = {str(s.pk): s.unit_type for s in Service.objects.all()}
    return render(request, 'proposals/form.html', {
        'form': form,
        'service_units_json': json.dumps(service_units),
    })


@login_required
@require_http_methods(['GET', 'POST'])
def confirm_view(request):
    """Шаг 2: подтверждение и скачивание PDF."""
    data = request.session.get('proposal_data')
    if not data:
        return redirect('proposals:form')

    if request.method == 'POST':
        if request.POST.get('save_draft'):
            _save_tkp_record(data, status=TKPRecord.STATUS_DRAFT, user=request.user)
            messages.success(request, 'Черновик ТКП сохранён в перечень.')
            return redirect('proposals:table')
        try:
            base_name = _generate_and_save_files(data)
        except Exception as e:
            messages.error(request, f'Ошибка генерации: {e}')
            return redirect('proposals:confirm')
        if base_name:
            _save_tkp_record(data, user=request.user)
            draft_id = request.session.pop('tkp_draft_id', None)
            if draft_id:
                TKPRecord.objects.filter(pk=draft_id).delete()
            request.session['tkp_download_base'] = base_name
            return redirect('proposals:download_success')
        messages.error(
            request,
            'Ошибка генерации. Проверьте, что шаблоны .docx есть в папке templates_docx/'
        )
        return redirect('proposals:confirm')

    from datetime import datetime
    date_display = datetime.strptime(data['date'], '%Y-%m-%d').strftime('%d.%m.%Y')
    price_display = _format_price(Decimal(data['price']))

    context = {
        'date': date_display,
        'service_name': data['service_name'],
        'city': data['city'],
        'price': price_display,
        'price_raw': data['price'],
        'client': data['client'],
        'room': data.get('room', ''),
        'srok': data.get('srok', ''),
        'text': data['text'],
        's': data['s'],
    }
    return render(request, 'proposals/confirm.html', context)


@login_required
@require_http_methods(['GET'])
def download_success_view(request):
    """Страница после формирования ТКП: ссылки на скачивание PDF и DOCX."""
    base_name = request.session.pop('tkp_download_base', None)
    if not base_name:
        return redirect('proposals:form')
    context = {
        'base_name': base_name,
        'pdf_name': f'{base_name}.pdf',
        'docx_name': f'{base_name}.docx',
    }
    return render(request, 'proposals/download_success.html', context)


@login_required
@require_http_methods(['GET'])
def download_file_view(request, file_type):
    """Отдача PDF или DOCX по base_name из GET-параметра (для скачивания после формирования)."""
    base_name = request.GET.get('f', '').strip()
    if not base_name or file_type not in ('pdf', 'docx'):
        raise Http404()
    # Допускаем буквы (латиница, кириллица), цифры, _, -, « », №
    if not re.match(r'^[a-zA-Z0-9_\-\u0400-\u04FF\u00AB\u00BB\u2116]+$', base_name):
        raise Http404()
    out_dir = Path(getattr(settings, 'TKP_OUTPUT_DIR', settings.BASE_DIR / 'TKP_output'))
    ext = 'pdf' if file_type == 'pdf' else 'docx'
    path = out_dir / f'{base_name}.{ext}'
    if not path.exists():
        raise Http404()
    return FileResponse(
        open(path, 'rb'),
        as_attachment=True,
        filename=path.name,
    )


def _parse_complex_rows(rows_data):
    """
    Валидация и разбор строк комплексного ТКП из JSON.
    Возвращает (list of dicts, error_message).
    Каждый dict: service_name, comment, unit, quantity, price_per_unit, total (Decimal).
    """
    if not rows_data:
        return [], 'Добавьте хотя бы одну строку'
    try:
        rows = json.loads(rows_data) if isinstance(rows_data, str) else rows_data
    except (json.JSONDecodeError, TypeError):
        return [], 'Неверный формат данных строк'
    if not isinstance(rows, list) or len(rows) == 0:
        return [], 'Добавьте хотя бы одну строку'
    result = []
    for i, r in enumerate(rows):
        if not isinstance(r, dict):
            continue
        name = (r.get('service_name') or '').strip()
        comment = (r.get('comment') or '').strip()
        srok = (r.get('srok') or '').strip()
        unit = (r.get('unit') or 'm2').strip() in ('piece', 'шт') and 'piece' or 'm2'
        try:
            qty = Decimal(str(r.get('quantity') or 0))
            price = Decimal(str(r.get('price_per_unit') or 0))
        except Exception:
            continue
        if qty < 0 or price < 0:
            continue
        total = (qty * price).quantize(Decimal('0.01'))
        result.append({
            'service_name': name or f'Позиция {i + 1}',
            'comment': comment,
            'srok': srok,
            'unit': unit,
            'quantity': qty,
            'price_per_unit': price,
            'total': total,
        })
    if not result:
        return [], 'Заполните хотя бы одну строку (компонент услуги, количество, цена)'
    return result, None


@login_required
@require_http_methods(['GET', 'POST'])
def complex_form_view(request):
    """Форма комплексного ТКП: дата, клиент, срок, таблица строк."""
    if request.method == 'POST':
        form = ComplexProposalForm(request.POST)
        rows_data = request.POST.get('rows_json', '')
        rows, row_error = _parse_complex_rows(rows_data)
        save_draft = request.POST.get('save_draft')

        if save_draft:
            # Черновик комплексного ТКП: минимум — дата и регион
            draft_errors = []
            date_val = form.data.get('date')
            if not date_val:
                draft_errors.append('Укажите дату.')
            else:
                try:
                    datetime.strptime(date_val, '%Y-%m-%d')
                except ValueError:
                    draft_errors.append('Некорректная дата.')
            region_id = form.data.get('region')
            region = None
            if region_id:
                try:
                    region = Region.objects.get(pk=region_id)
                except (Region.DoesNotExist, ValueError):
                    pass
            if not region:
                draft_errors.append('Выберите регион.')

            if not draft_errors and region:
                date_str = date_val
                client_val = (form.data.get('client') or '').strip()
                room_val = (form.data.get('room') or '').strip()
                text1_val = (form.data.get('text1') or '').strip()
                rows_serializable = []
                total_sum = Decimal(0)
                if not row_error and rows:
                    rows_serializable = [
                        {
                            'service_name': r['service_name'],
                            'comment': r.get('comment', ''),
                            'srok': r.get('srok', ''),
                            'unit': r['unit'],
                            'quantity': str(r['quantity']),
                            'price_per_unit': str(r['price_per_unit']),
                            'total': str(r['total']),
                        }
                        for r in rows
                    ]
                    total_sum = sum(Decimal(str(r['total'])) for r in rows)
                data = {
                    'date': date_str,
                    'client': client_val,
                    'region_id': region.pk,
                    'region_name': region.name,
                    'room': room_val,
                    'rows': rows_serializable,
                    'text1': text1_val,
                }
                request.session['complex_proposal_data'] = data
                _save_complex_tkp_record(data, status=TKPRecord.STATUS_DRAFT, user=request.user)
                messages.success(request, 'Черновик комплексного ТКП сохранён в перечень.')
                return redirect('proposals:table')
            for err in draft_errors:
                messages.error(request, err)
        elif form.is_valid() and not row_error:
            # Сессия сериализуется в JSON — храним числа как строки
            rows_serializable = [
                {
                    'service_name': r['service_name'],
                    'comment': r.get('comment', ''),
                    'srok': r.get('srok', ''),
                    'unit': r['unit'],
                    'quantity': str(r['quantity']),
                    'price_per_unit': str(r['price_per_unit']),
                    'total': str(r['total']),
                }
                for r in rows
            ]
            region = form.cleaned_data['region']
            data = {
                'date': form.cleaned_data['date'].strftime('%Y-%m-%d'),
                'client': (form.cleaned_data['client'] or '').strip(),
                'region_id': region.pk,
                'region_name': region.name,
                'room': (form.cleaned_data.get('room') or '').strip(),
                'rows': rows_serializable,
                'text1': (form.cleaned_data.get('text1') or '').strip(),
            }
            request.session['complex_proposal_data'] = data
            return redirect('proposals:complex_confirm')
        if row_error and not save_draft:
            messages.error(request, row_error)
    else:
        draft_id = request.GET.get('draft_id')
        initial = None
        if draft_id:
            try:
                draft = TKPRecord.objects.get(pk=draft_id, status=TKPRecord.STATUS_DRAFT)
                if draft.service == 'Комплексное ТКП':
                    initial = {
                        'date': draft.date,
                        'client': draft.client or '',
                        'room': draft.room or '',
                        'text1': draft.text or '',
                    }
                    request.session['complex_draft_id'] = int(draft_id)
            except (TKPRecord.DoesNotExist, ValueError, TypeError):
                request.session.pop('complex_draft_id', None)
        else:
            request.session.pop('complex_draft_id', None)
        form = ComplexProposalForm(initial=initial) if initial else ComplexProposalForm()
    services_raw = list(Service.objects.order_by('order', 'name').values('id', 'name', 'unit_type', 'description'))
    services = []
    file_comments = _load_complex_service_comments_file()
    for s in services_raw:
        name = s['name']
        display_name = COMPLEX_SERVICE_DISPLAY_NAMES.get(name, name)
        saved_desc = (s.get('description') or '').strip()
        default_comment = (
            saved_desc
            or file_comments.get(name, '')
            or COMPLEX_SERVICE_DEFAULT_COMMENTS.get(name, '')
        )
        services.append({
            'id': s['id'],
            'name': name,
            'display_name': display_name,
            'unit_type': s['unit_type'],
            'default_comment': default_comment,
        })
    prices_qs = RegionServicePrice.objects.select_related('region', 'service').all()
    region_prices = {}
    for rsp in prices_qs:
        rid, sid = str(rsp.region_id), str(rsp.service_id)
        if rid not in region_prices:
            region_prices[rid] = {}
        region_prices[rid][sid] = str(rsp.unit_price)
    service_comments = {str(s['id']): s['default_comment'] for s in services}
    context = {
        'form': form,
        'services': services,
        'srok_choices': SROK_CHOICES,
        'services_json': json.dumps(services, ensure_ascii=False),
        'region_prices_json': json.dumps(region_prices, ensure_ascii=False),
        'service_comments_json': json.dumps(service_comments, ensure_ascii=False),
    }
    return render(request, 'proposals/complex_form.html', context)


@login_required
@require_http_methods(['GET', 'POST'])
def complex_confirm_view(request):
    """Подтверждение комплексного ТКП и генерация docx/pdf."""
    data = request.session.get('complex_proposal_data')
    if not data:
        return redirect('proposals:complex_form')

    if request.method == 'POST':
        if request.POST.get('save_draft'):
            _save_complex_tkp_record(data, status=TKPRecord.STATUS_DRAFT, user=request.user)
            messages.success(request, 'Черновик комплексного ТКП сохранён в перечень.')
            return redirect('proposals:table')
        try:
            base_name = _generate_complex_and_save_files(data)
        except Exception as e:
            messages.error(request, f'Ошибка генерации: {e}')
            return redirect('proposals:complex_confirm')
        if base_name:
            _save_complex_tkp_record(data, user=request.user)
            draft_id = request.session.pop('complex_draft_id', None)
            if draft_id:
                TKPRecord.objects.filter(pk=draft_id).delete()
            request.session['tkp_download_base'] = base_name
            return redirect('proposals:download_success')
        messages.error(request, 'Ошибка генерации. Проверьте шаблон в templates_docx/')
        return redirect('proposals:complex_confirm')

    date_display = datetime.strptime(data['date'], '%Y-%m-%d').strftime('%d.%m.%Y')
    total_sum = sum(Decimal(str(r['total'])) for r in data['rows'])
    rows_display = []
    for i, r in enumerate(data['rows'], 1):
        comment = (r.get('comment') or '').replace(COMPLEX_COMMENT_LINE_BREAK_MARKER, '\n')
        srok = r.get('srok', '')
        if srok:
            comment = comment.rstrip() + '\nСрок разработки - ' + srok
        rows_display.append({
            'num': i,
            'service_name': r['service_name'],
            'comment': comment,
            'unit_display': UNIT_DISPLAY.get(r['unit'], r['unit']),
            'quantity': r['quantity'],
            'price_per_unit': r['price_per_unit'],
            'total': _format_price(Decimal(str(r['total']))),
        })
    context = {
        'date': date_display,
        'client': data['client'],
        'region_name': data.get('region_name', ''),
        'room': data.get('room', ''),
        'rows': rows_display,
        'total_sum': _format_price(total_sum),
        'text1': data.get('text1', ''),
    }
    return render(request, 'proposals/complex_confirm.html', context)


TKP_TABLE_PLACEHOLDER = '___TKP_TABLE_INSERT___'


def _set_table_borders(table, sz='4', val='single', color='000000'):
    """Включает границы таблицы через tblBorders (вся таблица)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = tblBorders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            tblBorders.append(el)
        el.set(qn('w:val'), val)
        el.set(qn('w:sz'), sz)
        el.set(qn('w:color'), color)


def _set_cell_font(cell, font_name='Montserrat'):
    """Устанавливает шрифт для всех run в ячейке."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name


def _set_cell_text_with_breaks_and_font(cell, text, font_name='Montserrat'):
    """Заполняет ячейку текстом с сохранением переносов строк, шрифт Montserrat."""
    from docx.enum.text import WD_BREAK
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    if not text:
        _set_cell_font(cell, font_name)
        return
    lines = text.split('\n')
    for i, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = font_name
        if i < len(lines) - 1:
            run.add_break(WD_BREAK.LINE)
    _set_cell_font(cell, font_name)


def _set_cell_comment_with_srok(cell, comment, srok, font_name='Montserrat'):
    """Ячейка «Комментарий»: текст с переносами (Montserrat), затем с новой строки жирная «Срок разработки - {srok}».
    Переносы: символ COMPLEX_COMMENT_LINE_BREAK_MARKER и реальные \\n заменяются на новую строку."""
    from docx.enum.text import WD_BREAK
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    comment = (comment or '').strip().replace(COMPLEX_COMMENT_LINE_BREAK_MARKER, '\n')
    srok = (srok or '').strip()
    if comment:
        lines = comment.split('\n')
        for i, line in enumerate(lines):
            run = paragraph.add_run(line)
            run.font.name = font_name
            if i < len(lines) - 1:
                run.add_break(WD_BREAK.LINE)
    if srok:
        if comment:
            run = paragraph.add_run()
            run.add_break(WD_BREAK.LINE)
        run = paragraph.add_run('Срок разработки - ' + srok)
        run.font.name = font_name
        run.bold = True
    if not comment and not srok:
        _set_cell_font(cell, font_name)
    else:
        _set_cell_font(cell, font_name)


def _build_complex_table_document(rows_ctx, total_sum_formatted):
    """Создаёт Document с одной таблицей позиций (для вставки в основной docx). Без колонки №.
    Ширины задаются и в ячейках (Word), и в tblGrid (LibreOffice). Шрифт таблицы — Montserrat."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Inches
    doc = Document()
    table = doc.add_table(rows=len(rows_ctx) + 2, cols=6)  # заголовок + строки + итого
    _set_table_borders(table)
    table.autofit = False
    if hasattr(table, 'allow_autofit'):
        table.allow_autofit = False
    table.width = Inches(6.5)
    # Ширины: Компонент 1.5", Комментарий 35% (2.28"), остальные фиксированы
    col_widths_inches = (1.5, 2.28, 0.55, 0.6, 0.85, 0.9)
    col_widths = tuple(Inches(w) for w in col_widths_inches)
    # Word: ширина на каждой ячейке
    for row in table.rows:
        for col_idx, width in enumerate(col_widths):
            row.cells[col_idx].width = width
    # LibreOffice (прод): ширина в tblGrid/gridCol (w:w в twips, 1" = 1440 twips)
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is not None:
        grid_cols = tbl_grid.findall(qn('w:gridCol'))
        for col_idx, w_inch in enumerate(col_widths_inches):
            if col_idx < len(grid_cols):
                grid_cols[col_idx].set(qn('w:w'), str(int(round(w_inch * 1440))))
    header = table.rows[0].cells
    headers_text = ('Компонент услуги', 'Комментарий', 'Ед. изм.', 'Количество', 'Цена за ед.', 'Стоимость')
    for i, text in enumerate(headers_text):
        header[i].text = text
        for run in header[i].paragraphs[0].runs:
            run.bold = True
            run.font.name = 'Montserrat'
    for i, r in enumerate(rows_ctx, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = r['service_name']
        _set_cell_comment_with_srok(row_cells[1], r.get('comment', ''), r.get('srok', ''))
        row_cells[2].text = r['unit_display']
        row_cells[3].text = r['quantity']
        row_cells[4].text = r['price_per_unit']
        row_cells[5].text = r['total_formatted']
        for col_idx in (0, 2, 3, 4, 5):
            _set_cell_font(row_cells[col_idx])
    last = table.rows[len(rows_ctx) + 1].cells
    last[0].merge(last[5])
    last[0].text = f"Итого: {total_sum_formatted} ₽"
    _set_cell_font(last[0])
    return doc


def _insert_table_into_docx(docx_path, table_doc, placeholder=None):
    """Находит в docx абзац с плейсхолдером и заменяет его на таблицу."""
    from copy import deepcopy
    from docx import Document
    if placeholder is None:
        placeholder = TKP_TABLE_PLACEHOLDER
    doc = Document(str(docx_path))
    for p in doc.paragraphs:
        if placeholder in p.text:
            table = table_doc.tables[0]
            table_elem = deepcopy(table._tbl)
            p._p.addnext(table_elem)
            p._p.getparent().remove(p._p)
            doc.save(str(docx_path))
            return
    # если абзац не найден, таблица не вставляется (шаблон без плейсхолдера)


def _generate_complex_and_save_files(data):
    """Генерация docx по шаблону комплексного ТКП, конвертация в PDF, сохранение. Возвращает base_name."""
    templates_dir = getattr(settings, 'TEMPLATES_DOCX_DIR', Path(settings.BASE_DIR) / 'templates_docx')
    template_path = templates_dir / COMPLEX_TEMPLATE_NAME
    if not template_path.exists():
        return None

    date_obj = datetime.strptime(data['date'], '%Y-%m-%d').date()
    date_display = date_obj.strftime('%d.%m.%Y')
    seq = _get_next_seq_for_date(date_obj)
    number = f'{date_obj:%d%m%Y}_{seq}'

    total_sum = sum(Decimal(str(r['total'])) for r in data['rows'])
    total_sum_formatted = _format_price(total_sum)
    rows_ctx = []
    for i, r in enumerate(data['rows'], 1):
        rows_ctx.append({
            'num': i,
            'service_name': r['service_name'],
            'comment': r.get('comment', ''),
            'srok': r.get('srok', ''),
            'unit_display': UNIT_DISPLAY.get(r['unit'], r['unit']),
            'quantity': str(r['quantity']),
            'price_per_unit': _format_price(Decimal(str(r['price_per_unit']))),
            'total_formatted': _format_price(Decimal(str(r['total']))),
        })

    default_row = {
        'num': '', 'service_name': '', 'comment': '', 'unit_display': '',
        'quantity': '', 'price_per_unit': '', 'total_formatted': '',
    }
    context = {
        'date': date_display,
        'client': data['client'],
        'room': data.get('room', ''),
        'srok': data.get('srok', ''),
        'number': number,
        'total_sum_formatted': total_sum_formatted,
        'row': default_row,
        'rows': rows_ctx,
        'rows_table_placeholder': TKP_TABLE_PLACEHOLDER,
        'text1': data.get('text1') or '',
    }

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir = Path(tmpdir)
        docx_path = tmpdir / 'tkp.docx'
        pdf_path = tmpdir / 'tkp.pdf'
        doc = DocxTemplate(str(template_path))
        doc.render(context)
        doc.save(str(docx_path))
        table_doc = _build_complex_table_document(rows_ctx, total_sum_formatted)
        _insert_table_into_docx(docx_path, table_doc)
        _convert_docx_to_pdf(docx_path, pdf_path.parent)
        client_safe = _sanitize_filename(data.get('client') or '')
        base_name = f'{client_safe}_{date_obj:%d%m%Y}_{seq}'
        out_dir = Path(getattr(settings, 'TKP_OUTPUT_DIR', settings.BASE_DIR / 'TKP_output'))
        out_dir.mkdir(parents=True, exist_ok=True)
        shutil.copy2(pdf_path, out_dir / f'{base_name}.pdf')
        shutil.copy2(docx_path, out_dir / f'{base_name}.docx')
        return base_name


def _serialize_complex_rows_for_storage(rows):
    """Сериализация строк комплексного ТКП для сохранения в rows_json."""
    if not rows:
        return None
    return [
        {
            'service_name': r.get('service_name', ''),
            'comment': r.get('comment', ''),
            'srok': r.get('srok', ''),
            'unit': r.get('unit', ''),
            'quantity': str(r.get('quantity', '')),
            'price_per_unit': str(r.get('price_per_unit', '')),
            'total': str(r.get('total', '')),
        }
        for r in rows
    ]


def _save_complex_tkp_record(data, status=None, user=None):
    """Сохранение записи о сформированном комплексном ТКП (status по умолчанию — итоговый)."""
    date_obj = datetime.strptime(data['date'], '%Y-%m-%d').date()
    rows = data.get('rows', [])
    total_sum = sum(Decimal(str(r['total'])) for r in rows)
    if status == TKPRecord.STATUS_DRAFT:
        seq = _get_next_draft_seq_for_date(date_obj)
        number = _generate_draft_number(date_obj, seq)
    else:
        seq = _get_next_seq_for_date(date_obj)
        number = _generate_doc_number(data.get('client') or '', date_obj, seq)
    rows_json = _serialize_complex_rows_for_storage(rows)
    TKPRecord.objects.create(
        date=date_obj,
        number=number,
        client=data.get('client') or '',
        service='Комплексное ТКП',
        sum_total=total_sum,
        room=data.get('room') or '',
        s='',
        text=data.get('text1') or '',
        status=status or TKPRecord.STATUS_FINAL,
        created_by=user,
        rows_json=rows_json,
    )


def _delete_tkp_files(base_name):
    """Удаление файлов PDF и DOCX из TKP_output по base_name (номеру документа)."""
    out_dir = Path(getattr(settings, 'TKP_OUTPUT_DIR', settings.BASE_DIR / 'TKP_output'))
    for ext in ('pdf', 'docx'):
        path = out_dir / f'{base_name}.{ext}'
        if path.exists():
            try:
                path.unlink()
            except OSError:
                pass


@login_required
@require_http_methods(['GET', 'POST'])
def table_view(request):
    """Страница перечня сформированных ТКП с фильтрами по каждому столбцу; удаление записей."""
    if request.method == 'POST':
        q = request.GET.urlencode()
        url = reverse('proposals:table') + ('?' + q if q else '')
        bulk_ids = request.POST.getlist('ids')
        bulk_action = (request.POST.get('bulk_action') or '').strip()

        if bulk_action == 'delete' and bulk_ids:
            deleted = 0
            for pk in bulk_ids:
                try:
                    rec = TKPRecord.objects.get(pk=pk)
                    base_name = rec.number
                    rec.delete()
                    if base_name:
                        _delete_tkp_files(base_name)
                    deleted += 1
                except (TKPRecord.DoesNotExist, ValueError, TypeError):
                    pass
            if deleted:
                messages.success(request, f'Удалено записей: {deleted}.')
            return redirect(url)

        if bulk_action == 'copy' and bulk_ids:
            copied = 0
            for pk in bulk_ids:
                try:
                    rec = TKPRecord.objects.get(pk=pk)
                    if rec.status == TKPRecord.STATUS_DRAFT:
                        seq = _get_next_draft_seq_for_date(rec.date)
                        new_number = _generate_draft_number(rec.date, seq)
                    else:
                        seq = _get_next_seq_for_date(rec.date)
                        new_number = _generate_doc_number(rec.client or '', rec.date, seq)
                    TKPRecord.objects.create(
                        date=rec.date,
                        number=new_number,
                        client=rec.client or '',
                        service=rec.service or '',
                        sum_total=rec.sum_total or Decimal(0),
                        room=rec.room or '',
                        s=rec.s or '',
                        text=rec.text or '',
                        status=rec.status,
                        rows_json=rec.rows_json,
                        created_by=request.user,
                    )
                    copied += 1
                except (TKPRecord.DoesNotExist, ValueError, TypeError):
                    pass
            if copied:
                messages.success(request, f'Скопировано записей: {copied}.')
            return redirect(url)

        delete_id = request.POST.get('delete_id', '').strip()
        if delete_id:
            try:
                rec = TKPRecord.objects.get(pk=delete_id)
                base_name = rec.number
                rec.delete()
                _delete_tkp_files(base_name)
                messages.success(request, 'Запись удалена.')
            except TKPRecord.DoesNotExist:
                messages.error(request, 'Запись не найдена.')
            return redirect(url)
    records = TKPRecord.objects.select_related('created_by').all()
    date_from = request.GET.get('date_from', '').strip()
    date_to = request.GET.get('date_to', '').strip()
    number = request.GET.get('number', '').strip()
    client = request.GET.get('client', '').strip()
    service = request.GET.get('service', '').strip()
    sum_min = request.GET.get('sum_min', '').strip()
    sum_max = request.GET.get('sum_max', '').strip()
    if date_from:
        try:
            d = datetime.strptime(date_from, '%Y-%m-%d').date()
            records = records.filter(date__gte=d)
        except ValueError:
            pass
    if date_to:
        try:
            d = datetime.strptime(date_to, '%Y-%m-%d').date()
            records = records.filter(date__lte=d)
        except ValueError:
            pass
    if number:
        records = records.filter(number__icontains=number)
    if client:
        records = records.filter(client__icontains=client)
    if service:
        records = records.filter(service=service)
    if sum_min:
        try:
            records = records.filter(sum_total__gte=Decimal(sum_min.replace(',', '.')))
        except (ValueError, Exception):
            pass
    if sum_max:
        try:
            records = records.filter(sum_total__lte=Decimal(sum_max.replace(',', '.')))
        except (ValueError, Exception):
            pass
    services = list(TKPRecord.objects.values_list('service', flat=True).distinct().order_by('service'))
    contract_template_by_service = SERVICE_TO_CONTRACT_TEMPLATE
    contract_by_tkp = {
        c.tkp_id: c
        for c in ContractRecord.objects.filter(
            status=ContractRecord.STATUS_FINAL, tkp__isnull=False
        ).select_related('tkp')
    }
    contract_draft_by_tkp = {
        c.tkp_id: c
        for c in ContractRecord.objects.filter(
            status=ContractRecord.STATUS_DRAFT, tkp__isnull=False
        ).select_related('tkp')
    }
    contract_template_by_tkp = {}
    for r in records:
        if r.service == 'Комплексное ТКП' and r.rows_json:
            tpl = get_contract_template_for_complex_tkp(r.rows_json)
            # Показываем кнопку для любого комплексного ТКП с сохранёнными строками;
            # при отсутствии шаблона contract_form_view покажет сообщение
            contract_template_by_tkp[r.pk] = tpl if tpl else True
    context = {
        'records': records,
        'filters': {
            'date_from': date_from,
            'date_to': date_to,
            'number': number,
            'client': client,
            'service': service,
            'sum_min': sum_min,
            'sum_max': sum_max,
        },
        'services_list': services,
        'contract_template_by_service': contract_template_by_service,
        'contract_template_by_tkp': contract_template_by_tkp,
        'contract_by_tkp': contract_by_tkp,
        'contract_draft_by_tkp': contract_draft_by_tkp,
    }
    return render(request, 'proposals/table.html', context)


@login_required
@require_http_methods(['GET', 'POST'])
def contract_form_view(request, tkp_id):
    """Форма реквизитов договора по выбранному ТКП; генерация docx."""
    try:
        tkp = TKPRecord.objects.get(pk=tkp_id)
    except TKPRecord.DoesNotExist:
        messages.error(request, 'Запись ТКП не найдена.')
        return redirect('proposals:table')

    COMPLEX_CONTRACT_TEMPLATE_05 = '05_Договор_Контент_Навигация.docx'
    COMPLEX_CONTRACT_TEMPLATE_08 = '08_Договор_ДПФ_Благоустройство.docx'

    if tkp.service == 'Комплексное ТКП':
        is_complex_contract = True
        # Шаблон выбирается пользователем в форме (поле «Комплексный договор»), по умолчанию — 05
        contract_template_file = None  # задаётся ниже из initial или из form
    else:
        contract_template_file = SERVICE_TO_CONTRACT_TEMPLATE.get(tkp.service)
        if not contract_template_file:
            messages.error(request, f'Для услуги «{tkp.service}» формирование договора не предусмотрено.')
            return redirect('proposals:table')
        is_complex_contract = False

    templates_dir = getattr(settings, 'TEMPLATES_DOCX_DIR', Path(settings.BASE_DIR) / 'templates_docx')

    # Предзаполнение: из черновика (GET contract_draft_id) или из ТКП и карточки контрагента
    contract_draft_id = None
    draft_record = None
    if request.method != 'POST':
        draft_id_param = request.GET.get('contract_draft_id')
        if draft_id_param:
            try:
                draft_record = ContractRecord.objects.filter(
                    pk=draft_id_param,
                    status=ContractRecord.STATUS_DRAFT,
                    tkp_id=tkp_id,
                ).select_related('counterparty').first()
                if draft_record:
                    contract_draft_id = draft_record.pk
            except (ValueError, TypeError):
                pass

    date_str = tkp.date.strftime('%Y-%m-%d')
    next_contract_number = f'{tkp.date:%d%m%Y}_{_get_next_contract_seq_for_date(tkp.date)}'
    initial = {
        'contract_number': next_contract_number,
        'date': date_str,
        'price': tkp.sum_total,
        'payment_terms': DEFAULT_PAYMENT_TERMS,
        'room': tkp.room or '',
        's': tkp.s or '',
    }
    if is_complex_contract:
        initial['complex_contract_type'] = COMPLEX_CONTRACT_TEMPLATE_05
    if draft_record:
        initial['contract_number'] = draft_record.number
        initial['date'] = draft_record.date.strftime('%Y-%m-%d')
        initial['price'] = draft_record.sum_total

    # Контрагент: из черновика, из GET (возврат из справочника) или по совпадению с клиентом ТКП
    cp_for_initial = None
    if request.method != 'POST':
        if draft_record and draft_record.counterparty_id:
            cp_for_initial = draft_record.counterparty
        if not cp_for_initial:
            cp_id = request.GET.get('counterparty_id')
            if cp_id:
                try:
                    cp_for_initial = Counterparty.objects.get(pk=cp_id)
                except (Counterparty.DoesNotExist, ValueError):
                    pass
    if cp_for_initial:
        initial['counterparty'] = cp_for_initial.pk
        initial['customer_name'] = cp_for_initial.name or ''  # Наименование заказчика из карточки контрагента
        initial['customer_represented_by'] = _director_genitive(cp_for_initial.director or '')
        initial['customer_represented_by_nominative'] = cp_for_initial.director or ''  # им.п. для второго вхождения
        initial['name'] = cp_for_initial.name or ''
        initial['address'] = cp_for_initial.address or ''
        initial['inn'] = cp_for_initial.inn or ''
        initial['kpp'] = cp_for_initial.kpp or ''
        initial['ogrn'] = cp_for_initial.ogrn or ''
        initial['account'] = cp_for_initial.account or ''
        initial['bank'] = cp_for_initial.bank or ''
        initial['bik'] = cp_for_initial.bik or ''
        initial['kor_account'] = cp_for_initial.kor_account or ''
        initial['email'] = cp_for_initial.email or ''
    else:
        initial['customer_name'] = tkp.client or ''

    # Для комплексного ТКП шаблон берём из выбора пользователя (по умолчанию 05)
    if is_complex_contract:
        contract_template_file = initial.get('complex_contract_type') or COMPLEX_CONTRACT_TEMPLATE_05
    template_path = templates_dir / CONTRACT_TEMPLATES_SUBDIR / contract_template_file
    if not template_path.exists():
        messages.error(request, f'Шаблон договора не найден: {contract_template_file}')
        return redirect('proposals:table')

    if request.method == 'POST':
        form = ContractForm(request.POST)
        # Для комплексного ТКП шаблон берём из выбора в форме
        if is_complex_contract:
            ct = (request.POST.get('complex_contract_type') or '').strip()
            if ct in (COMPLEX_CONTRACT_TEMPLATE_05, COMPLEX_CONTRACT_TEMPLATE_08):
                contract_template_file = ct
            else:
                contract_template_file = COMPLEX_CONTRACT_TEMPLATE_05
            template_path = templates_dir / CONTRACT_TEMPLATES_SUBDIR / contract_template_file
        save_draft = request.POST.get('save_draft')
        if save_draft:
            draft_errors = []
            cp_id = form.data.get('counterparty')
            date_val = form.data.get('date')
            number_val = (form.data.get('contract_number') or '').strip()
            if not cp_id:
                draft_errors.append('Выберите контрагента.')
            else:
                try:
                    cp = Counterparty.objects.get(pk=cp_id)
                except (Counterparty.DoesNotExist, ValueError):
                    cp = None
                    draft_errors.append('Выберите контрагента.')
            if not date_val:
                draft_errors.append('Укажите дату договора.')
            else:
                try:
                    date_obj = datetime.strptime(date_val, '%Y-%m-%d').date()
                except ValueError:
                    date_obj = None
                    draft_errors.append('Некорректная дата.')
            if not number_val:
                draft_errors.append('Укажите номер договора.')
            else:
                number_qs = ContractRecord.objects.filter(number=number_val)
                draft_pk = request.POST.get('contract_draft_id')
                if draft_pk:
                    try:
                        existing_draft = ContractRecord.objects.get(
                            pk=draft_pk,
                            status=ContractRecord.STATUS_DRAFT,
                            tkp=tkp,
                        )
                        number_qs = number_qs.exclude(pk=existing_draft.pk)
                    except (ContractRecord.DoesNotExist, ValueError, TypeError):
                        pass
                if date_obj and number_qs.exists():
                    draft_errors.append('Договор с таким номером уже существует.')
            if not draft_errors and cp and date_obj and number_val:
                draft_pk = request.POST.get('contract_draft_id')
                try:
                    existing_draft = ContractRecord.objects.get(
                        pk=draft_pk,
                        status=ContractRecord.STATUS_DRAFT,
                        tkp=tkp,
                    )
                except (ContractRecord.DoesNotExist, ValueError, TypeError):
                    existing_draft = None
                if existing_draft:
                    price_raw = form.data.get('price')
                    try:
                        sum_total_val = Decimal(str(price_raw).replace(',', '.')) if price_raw else (tkp.sum_total or Decimal(0))
                    except (ValueError, TypeError, Exception):
                        sum_total_val = tkp.sum_total or Decimal(0)
                    existing_draft.date = date_obj
                    existing_draft.number = number_val
                    existing_draft.counterparty = cp
                    existing_draft.client = tkp.client or ''
                    existing_draft.service = tkp.service or ''
                    existing_draft.sum_total = sum_total_val
                    existing_draft.save()
                    messages.success(request, 'Черновик договора обновлён.')
                else:
                    ContractRecord.objects.create(
                        date=date_obj,
                        number=number_val,
                        status=ContractRecord.STATUS_DRAFT,
                        tkp=tkp,
                        counterparty=cp,
                        client=tkp.client or '',
                        service=tkp.service or '',
                        sum_total=tkp.sum_total or Decimal(0),
                        created_by=request.user,
                    )
                    messages.success(request, 'Черновик договора сохранён.')
                return redirect('proposals:table')
            for err in draft_errors:
                messages.error(request, err)
            contract_draft_id = request.POST.get('contract_draft_id')
        elif request.POST.get('preview_edit') and form.is_valid():
            if is_complex_contract and not form.cleaned_data.get('complex_contract_type'):
                form.add_error('complex_contract_type', 'Выберите тип комплексного договора.')
            if not form.errors:
                cd = form.cleaned_data
                if is_complex_contract:
                    contract_template_file = cd['complex_contract_type']
                    template_path = templates_dir / CONTRACT_TEMPLATES_SUBDIR / contract_template_file
                cp = cd['counterparty']
                date_obj = cd['date']
                price_val = cd['price']
                seq = _get_next_contract_seq_for_date(date_obj)
                contract_number = f'{date_obj:%d%m%Y}_{seq}'
                customer_name = (cd.get('customer_name') or '').strip() or (cp.name or '')
                customer_represented_by = (cd.get('customer_represented_by') or '').strip()
                if not customer_represented_by:
                    customer_represented_by = _director_genitive(cp.director or '')
                customer_represented_by_nominative = (cd.get('customer_represented_by_nominative') or '').strip() or (cp.director or '')
                payment_terms = (cd.get('payment_terms') or '').strip() or DEFAULT_PAYMENT_TERMS
                customer_in_person_raw = (cd.get('customer_in_person') or '').strip()
                ctx = {
                    'contract_number': contract_number,
                    'number': contract_number,
                    'date': date_obj.strftime('%d.%m.%Y'),
                    'customer_name': customer_name,
                    'customer_represented_by': customer_represented_by,
                    'customer_represented_by_nominative': customer_represented_by_nominative,
                    'customer_in_person': customer_in_person_raw,
                    'dolznost': _dolznost_from_customer_in_person(customer_in_person_raw),
                    'acting_on_basis': (cd.get('acting_on_basis') or '').strip(),
                    'work_completion_period': (cd.get('work_completion_period') or '').strip(),
                    'period_starts_from': (cd.get('period_starts_from') or '').strip(),
                    'price': _format_price(price_val),
                    'payment_terms': payment_terms,
                    'name': cd.get('name') or cp.name or '',
                    'address': cd.get('address') or cp.address or '',
                    'inn': cd.get('inn') or cp.inn or '',
                    'kpp': cd.get('kpp') or cp.kpp or '',
                    'ogrn': cd.get('ogrn') or cp.ogrn or '',
                    'account': cd.get('account') or cp.account or '',
                    'bank': cd.get('bank') or cp.bank or '',
                    'bik': cd.get('bik') or cp.bik or '',
                    'kor_account': cd.get('kor_account') or cp.kor_account or '',
                    'email': cd.get('email') or cp.email or '',
                    'room': cd.get('room') or tkp.room or '',
                    's': cd.get('s') or tkp.s or '',
                }
                doc = DocxTemplate(str(template_path))
                doc.render(ctx)
                _set_contract_doc_font_times_new_roman(doc)
                if is_complex_contract and tkp.rows_json:
                    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                        doc.save(tmp.name)
                        rows_ctx, total_fmt = _complex_rows_json_to_ctx(tkp.rows_json)
                        table_doc = _build_complex_table_document(rows_ctx, total_fmt)
                        _insert_table_into_docx(tmp.name, table_doc, CONTRACT_SPEC_TABLE_PLACEHOLDER)
                        with open(tmp.name, 'rb') as f:
                            buf = io.BytesIO(f.read())
                        try:
                            os.unlink(tmp.name)
                        except OSError:
                            pass
                else:
                    buf = io.BytesIO()
                    doc.save(buf)
                buf.seek(0)
                try:
                    result = mammoth.convert_to_html(buf)
                    html_content = result.value or '<p>Не удалось преобразовать документ в HTML.</p>'
                except Exception:
                    html_content = '<p>Ошибка преобразования черновика в HTML. Проверьте шаблон договора.</p>'
                request.session['contract_editor_html'] = html_content
                request.session['contract_editor_tkp_id'] = tkp_id
                request.session['contract_editor_contract_number'] = contract_number
                request.session['contract_editor_date'] = date_obj.strftime('%Y-%m-%d')
                request.session['contract_editor_price'] = str(price_val)
                request.session['contract_editor_counterparty_id'] = cp.pk
                request.session['contract_editor_draft_id'] = request.POST.get('contract_draft_id') or ''
                request.session['contract_editor_template_file'] = contract_template_file
                return redirect('proposals:contract_editor')
        elif form.is_valid():
            if is_complex_contract and not form.cleaned_data.get('complex_contract_type'):
                form.add_error('complex_contract_type', 'Выберите тип комплексного договора.')
            if not form.errors:
                cd = form.cleaned_data
                if is_complex_contract:
                    contract_template_file = cd['complex_contract_type']
                    template_path = templates_dir / CONTRACT_TEMPLATES_SUBDIR / contract_template_file
                cp = cd['counterparty']
                date_obj = cd['date']
                price_val = cd['price']
                seq = _get_next_contract_seq_for_date(date_obj)
                contract_number = f'{date_obj:%d%m%Y}_{seq}'
                customer_name = (cd.get('customer_name') or '').strip() or (cp.name or '')
                customer_represented_by = (cd.get('customer_represented_by') or '').strip()
                if not customer_represented_by:
                    customer_represented_by = _director_genitive(cp.director or '')
                customer_represented_by_nominative = (cd.get('customer_represented_by_nominative') or '').strip() or (cp.director or '')
                payment_terms = (cd.get('payment_terms') or '').strip() or DEFAULT_PAYMENT_TERMS
                customer_in_person_raw = (cd.get('customer_in_person') or '').strip()
                ctx = {
                    'contract_number': contract_number,
                    'number': contract_number,
                    'date': date_obj.strftime('%d.%m.%Y'),
                    'customer_name': customer_name,
                    'customer_represented_by': customer_represented_by,
                    'customer_represented_by_nominative': customer_represented_by_nominative,
                    'customer_in_person': customer_in_person_raw,
                    'dolznost': _dolznost_from_customer_in_person(customer_in_person_raw),
                    'acting_on_basis': (cd.get('acting_on_basis') or '').strip(),
                    'work_completion_period': (cd.get('work_completion_period') or '').strip(),
                    'period_starts_from': (cd.get('period_starts_from') or '').strip(),
                    'price': _format_price(price_val),
                    'payment_terms': payment_terms,
                    'name': cd.get('name') or cp.name or '',
                    'address': cd.get('address') or cp.address or '',
                    'inn': cd.get('inn') or cp.inn or '',
                    'kpp': cd.get('kpp') or cp.kpp or '',
                    'ogrn': cd.get('ogrn') or cp.ogrn or '',
                    'account': cd.get('account') or cp.account or '',
                    'bank': cd.get('bank') or cp.bank or '',
                    'bik': cd.get('bik') or cp.bik or '',
                    'kor_account': cd.get('kor_account') or cp.kor_account or '',
                    'email': cd.get('email') or cp.email or '',
                    'room': cd.get('room') or tkp.room or '',
                    's': cd.get('s') or tkp.s or '',
                }
                doc = DocxTemplate(str(template_path))
                doc.render(ctx)
                _set_contract_doc_font_times_new_roman(doc)
                out_format = (request.POST.get('format') or 'docx').strip().lower()
                if out_format != 'pdf':
                    out_format = 'docx'
                file_base = f'Дог_{contract_number}'
                out_dir = Path(getattr(settings, 'TKP_OUTPUT_DIR', settings.BASE_DIR / 'TKP_output'))
                out_dir.mkdir(parents=True, exist_ok=True)
                docx_path = out_dir / f'{file_base}.docx'
                pdf_path = out_dir / f'{file_base}.pdf'
                doc.save(str(docx_path))
                if is_complex_contract and tkp.rows_json:
                    rows_ctx, total_fmt = _complex_rows_json_to_ctx(tkp.rows_json)
                    table_doc = _build_complex_table_document(rows_ctx, total_fmt)
                    _insert_table_into_docx(str(docx_path), table_doc, CONTRACT_SPEC_TABLE_PLACEHOLDER)
                with tempfile.TemporaryDirectory() as tmpdir:
                    tmpdir = Path(tmpdir)
                    shutil.copy2(docx_path, tmpdir / 'contract.docx')
                    _convert_docx_to_pdf(tmpdir / 'contract.docx', tmpdir)
                    src = tmpdir / 'contract.pdf'
                    if not src.exists():
                        src = tmpdir / 'tkp.pdf'
                    if not src.exists():
                        src = next(tmpdir.glob('*.pdf'), None)
                    if src and src.exists():
                        shutil.copy2(src, pdf_path)
                ContractRecord.objects.create(
                    date=date_obj,
                    number=contract_number,
                    status=ContractRecord.STATUS_FINAL,
                    tkp=tkp,
                    counterparty=cp,
                    client=tkp.client or '',
                    service=tkp.service or '',
                    sum_total=price_val,
                    docx_file=file_base,
                    pdf_file=file_base,
                    contract_snapshot=ctx,
                    created_by=request.user,
                )
                if out_format == 'docx':
                    return FileResponse(
                        open(docx_path, 'rb'),
                        as_attachment=True,
                        filename=f'{file_base}.docx',
                    )
                if pdf_path.exists():
                    return FileResponse(
                        open(pdf_path, 'rb'),
                        as_attachment=True,
                        filename=f'{file_base}.pdf',
                    )
                messages.error(request, 'Не удалось сформировать PDF. Установите LibreOffice или docx2pdf.')
        # form errors: show form again
    else:
        form = ContractForm(initial=initial)

    counterparty_display_name = ''
    if request.method != 'POST' and cp_for_initial:
        counterparty_display_name = cp_for_initial.name or ''
    elif request.method == 'POST' and form.data.get('counterparty'):
        try:
            _cp = Counterparty.objects.get(pk=form.data.get('counterparty'))
            counterparty_display_name = _cp.name or ''
        except (Counterparty.DoesNotExist, ValueError, TypeError):
            pass

    return render(request, 'proposals/contract_form.html', {
        'form': form,
        'tkp': tkp,
        'contract_template_file': contract_template_file,
        'contract_draft_id': contract_draft_id,
        'counterparty_display_name': counterparty_display_name,
        'is_complex_contract': is_complex_contract,
    })


SESSION_KEY_EDITOR_HTML = 'contract_editor_html'
SESSION_KEY_EDITOR_TKP_ID = 'contract_editor_tkp_id'
SESSION_KEY_EDITOR_CONTRACT_NUMBER = 'contract_editor_contract_number'
SESSION_KEY_EDITOR_DATE = 'contract_editor_date'
SESSION_KEY_EDITOR_PRICE = 'contract_editor_price'
SESSION_KEY_EDITOR_COUNTERPARTY_ID = 'contract_editor_counterparty_id'
SESSION_KEY_EDITOR_DRAFT_ID = 'contract_editor_draft_id'
SESSION_KEY_EDITOR_TEMPLATE_FILE = 'contract_editor_template_file'


@login_required
@require_http_methods(['GET'])
def contract_editor_view(request):
    """Страница предпросмотра и редактирования черновика договора в CKEditor 5."""
    html_content = request.session.get(SESSION_KEY_EDITOR_HTML)
    if not html_content:
        messages.error(request, 'Сессия предпросмотра истекла или не найдена. Заполните реквизиты и нажмите «Предпросмотр и редактирование черновика».')
        return redirect('proposals:table')
    tkp_id = request.session.get(SESSION_KEY_EDITOR_TKP_ID)
    try:
        tkp = TKPRecord.objects.get(pk=tkp_id)
    except (TKPRecord.DoesNotExist, ValueError, TypeError):
        messages.error(request, 'Запись ТКП не найдена.')
        return redirect('proposals:table')
    return render(request, 'proposals/contract_editor.html', {
        'html_content': html_content,
        'tkp': tkp,
        'contract_number': request.session.get(SESSION_KEY_EDITOR_CONTRACT_NUMBER, ''),
        'contract_date': request.session.get(SESSION_KEY_EDITOR_DATE, ''),
        'contract_price': request.session.get(SESSION_KEY_EDITOR_PRICE, ''),
        'counterparty_id': request.session.get(SESSION_KEY_EDITOR_COUNTERPARTY_ID) or '',
        'contract_draft_id': request.session.get(SESSION_KEY_EDITOR_DRAFT_ID) or '',
    })


@login_required
@require_http_methods(['POST'])
def contract_save_from_editor_view(request):
    """Сохранение договора из отредактированного HTML (CKEditor): конвертация в DOCX и запись в реестр."""
    html_content = (request.POST.get('contract_editor_content') or '').strip()
    if not html_content:
        messages.error(request, 'Содержимое договора пусто.')
        return redirect('proposals:table')
    tkp_id = request.session.get(SESSION_KEY_EDITOR_TKP_ID)
    contract_number = (request.session.get(SESSION_KEY_EDITOR_CONTRACT_NUMBER) or '').strip()
    date_str = request.session.get(SESSION_KEY_EDITOR_DATE) or ''
    price_str = request.session.get(SESSION_KEY_EDITOR_PRICE) or '0'
    counterparty_id = request.session.get(SESSION_KEY_EDITOR_COUNTERPARTY_ID)
    if not tkp_id or not contract_number or not date_str:
        messages.error(request, 'Сессия предпросмотра истекла. Повторите ввод реквизитов и предпросмотр.')
        return redirect('proposals:table')
    try:
        tkp = TKPRecord.objects.get(pk=tkp_id)
    except (TKPRecord.DoesNotExist, ValueError, TypeError):
        messages.error(request, 'Запись ТКП не найдена.')
        return redirect('proposals:table')
    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d').date()
    except ValueError:
        messages.error(request, 'Некорректная дата.')
        return redirect('proposals:table')
    try:
        price_val = Decimal(price_str.replace(',', '.'))
    except (ValueError, TypeError):
        price_val = tkp.sum_total or Decimal(0)
    try:
        cp = Counterparty.objects.get(pk=counterparty_id)
    except (Counterparty.DoesNotExist, ValueError, TypeError):
        messages.error(request, 'Контрагент не найден.')
        return redirect('proposals:table')
    out_dir = Path(getattr(settings, 'TKP_OUTPUT_DIR', settings.BASE_DIR / 'TKP_output'))
    out_dir.mkdir(parents=True, exist_ok=True)
    file_base = f'Дог_{contract_number}'
    docx_path = out_dir / f'{file_base}.docx'

    templates_dir = Path(getattr(settings, 'TEMPLATES_DOCX_DIR', settings.BASE_DIR / 'templates_docx'))
    template_filename = request.session.get(SESSION_KEY_EDITOR_TEMPLATE_FILE)
    reference_doc = (templates_dir / CONTRACT_TEMPLATES_SUBDIR / template_filename) if template_filename else None
    extra_args = []
    if reference_doc and reference_doc.exists():
        extra_args.append('--reference-doc=' + str(reference_doc))

    def _html_to_docx():
        with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as f:
            f.write(html_content)
            html_path = f.name
        try:
            pypandoc.convert_file(html_path, 'docx', format='html', outputfile=str(docx_path), extra_args=extra_args)
        finally:
            try:
                os.unlink(html_path)
            except OSError:
                pass

    try:
        _html_to_docx()
    except OSError as e:
        err_msg = str(e)
        if 'pandoc' in err_msg.lower() and hasattr(pypandoc, 'download_pandoc'):
            try:
                pypandoc.download_pandoc()
                _html_to_docx()
            except Exception as e2:
                messages.error(request, f'Не удалось сформировать DOCX: {e2}. Установите Pandoc (https://pandoc.org) или используйте пакет pypandoc-binary.')
                return redirect('proposals:contract_editor')
        else:
            messages.error(request, f'Не удалось сформировать DOCX из текста: {e}. Установите Pandoc (https://pandoc.org) или установите пакет: pip install pypandoc-binary')
            return redirect('proposals:contract_editor')
    except Exception as e:
        messages.error(request, f'Не удалось сформировать DOCX из текста: {e}. Установите Pandoc (https://pandoc.org).')
        return redirect('proposals:contract_editor')
    if template_filename in ('05_Договор_Контент_Навигация.docx', '08_Договор_ДПФ_Благоустройство.docx') and tkp.rows_json:
        rows_ctx, total_fmt = _complex_rows_json_to_ctx(tkp.rows_json)
        table_doc = _build_complex_table_document(rows_ctx, total_fmt)
        _insert_table_into_docx(str(docx_path), table_doc, CONTRACT_SPEC_TABLE_PLACEHOLDER)
    ContractRecord.objects.create(
        date=date_obj,
        number=contract_number,
        status=ContractRecord.STATUS_FINAL,
        tkp=tkp,
        counterparty=cp,
        client=tkp.client or '',
        service=tkp.service or '',
        sum_total=price_val,
        docx_file=file_base,
        pdf_file=file_base,
        contract_snapshot={},
        created_by=request.user,
    )
    for key in (SESSION_KEY_EDITOR_HTML, SESSION_KEY_EDITOR_TKP_ID, SESSION_KEY_EDITOR_CONTRACT_NUMBER,
                SESSION_KEY_EDITOR_DATE, SESSION_KEY_EDITOR_PRICE, SESSION_KEY_EDITOR_COUNTERPARTY_ID,
                SESSION_KEY_EDITOR_DRAFT_ID, SESSION_KEY_EDITOR_TEMPLATE_FILE):
        request.session.pop(key, None)
    messages.success(request, 'Договор сохранён.')
    save_format = (request.POST.get('save_format') or 'docx').strip().lower()
    if save_format == 'pdf':
        pdf_path = out_dir / f'{file_base}.pdf'
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            _convert_docx_to_pdf(docx_path, tmpdir)
            src = tmpdir / f'{docx_path.stem}.pdf'  # LibreOffice: то же имя
            if not src.exists():
                src = tmpdir / 'tkp.pdf'  # docx2pdf
            if not src.exists():
                src = next(tmpdir.glob('*.pdf'), None)
            if src and src.exists():
                shutil.copy2(src, pdf_path)
        if pdf_path.exists():
            return FileResponse(
                open(pdf_path, 'rb'),
                as_attachment=True,
                filename=f'{file_base}.pdf',
            )
        messages.warning(request, 'PDF не сформирован. Скачан DOCX. Установите LibreOffice или docx2pdf для генерации PDF.')
    return FileResponse(
        open(docx_path, 'rb'),
        as_attachment=True,
        filename=f'{file_base}.docx',
    )


# Канбан: ключи колонок
KANBAN_COL_DRAFT = 'draft'
KANBAN_COL_FINAL = 'final'
KANBAN_COL_CONTRACT_DRAFT = 'contract_draft'
KANBAN_COL_CONTRACT_FINAL = 'contract_final'


def _kanban_computed_column(tkp, contract):
    """Вычисляет колонку канбана по статусу ТКП и договора."""
    if tkp.status == TKPRecord.STATUS_DRAFT:
        return KANBAN_COL_DRAFT
    if tkp.status != TKPRecord.STATUS_FINAL:
        return KANBAN_COL_DRAFT
    if contract is None:
        return KANBAN_COL_FINAL
    if contract.status == ContractRecord.STATUS_DRAFT:
        return KANBAN_COL_CONTRACT_DRAFT
    return KANBAN_COL_CONTRACT_FINAL


@login_required
@require_http_methods(['GET'])
def kanban_view(request):
    """Канбан-доска с переименованием колонок, пользовательскими колонками и позициями карточек."""
    tkp_list = list(TKPRecord.objects.order_by('-created_at'))
    contract_by_tkp = {
        c.tkp_id: c
        for c in ContractRecord.objects.filter(tkp__isnull=False).select_related('tkp')
    }
    default_titles = {
        KANBAN_COL_DRAFT: 'Черновик',
        KANBAN_COL_FINAL: 'Итоговый',
        KANBAN_COL_CONTRACT_DRAFT: 'Договор — черновик',
        KANBAN_COL_CONTRACT_FINAL: 'Договор сформирован',
    }
    title_overrides = {
        o.column_key: o.title
        for o in KanbanColumnTitleOverride.objects.filter(user=request.user)
    }
    custom_columns = list(
        KanbanColumnCustom.objects.filter(user=request.user).order_by('order', 'pk')
    )
    positions = {
        p.tkp_id: p.column_key
        for p in KanbanCardPosition.objects.filter(user=request.user)
    }
    tkp_ids = [t.pk for t in tkp_list]
    custom_fields_qs = KanbanCardField.objects.filter(
        user=request.user, tkp_id__in=tkp_ids
    ).order_by('tkp_id', 'order', 'pk')
    custom_fields_by_tkp = {}
    for f in custom_fields_qs:
        custom_fields_by_tkp.setdefault(f.tkp_id, []).append({'name': f.name, 'value': f.value or ''})

    all_column_keys = [KANBAN_COL_DRAFT, KANBAN_COL_FINAL, KANBAN_COL_CONTRACT_DRAFT, KANBAN_COL_CONTRACT_FINAL]
    columns = {key: {'title': title_overrides.get(key) or default_titles[key], 'cards': [], 'count': 0, 'sum': Decimal(0)} for key in all_column_keys}
    for cc in custom_columns:
        ckey = f'custom_{cc.pk}'
        columns[ckey] = {'title': cc.title, 'cards': [], 'count': 0, 'sum': Decimal(0), 'custom_id': cc.pk}

    for tkp in tkp_list:
        contract = contract_by_tkp.get(tkp.pk)
        col_key = positions.get(tkp.pk) or _kanban_computed_column(tkp, contract)
        if col_key not in columns:
            col_key = _kanban_computed_column(tkp, contract)
        service_display = 'Комплексное' if tkp.service == 'Комплексное ТКП' else tkp.service
        sum_val = tkp.sum_total or Decimal(0)
        card = {
            'tkp_id': tkp.pk,
            'tkp_number': tkp.number,
            'client': tkp.client or '—',
            'date': tkp.date,
            'date_display': tkp.date.strftime('%d.%m.%Y'),
            'service_display': service_display,
            'sum_total': sum_val,
            'sum_display': _format_price(sum_val),
            'contract_id': contract.pk if contract else None,
            'custom_fields': custom_fields_by_tkp.get(tkp.pk, []),
        }
        columns[col_key]['cards'].append(card)
        columns[col_key]['count'] += 1
        columns[col_key]['sum'] += sum_val

    default_order = [
        KANBAN_COL_DRAFT,
        KANBAN_COL_FINAL,
        KANBAN_COL_CONTRACT_DRAFT,
        KANBAN_COL_CONTRACT_FINAL,
    ] + [f'custom_{cc.pk}' for cc in custom_columns]
    try:
        board_order = KanbanBoardOrder.objects.get(user=request.user)
        saved_order = list(board_order.order) if isinstance(board_order.order, list) else []
    except KanbanBoardOrder.DoesNotExist:
        saved_order = []
    ordered_keys = [k for k in saved_order if k in columns] if saved_order else []
    for k in default_order:
        if k not in ordered_keys and k in columns:
            ordered_keys.append(k)

    def col_data(key):
        c = columns.get(key)
        if not c:
            return None
        return {
            'id': key,
            'title': c['title'],
            'cards': c['cards'],
            'count': c['count'],
            'sum': c['sum'],
            'sum_display': _format_price(c['sum']),
            'custom_id': c.get('custom_id'),
        }

    all_columns = [col_data(k) for k in ordered_keys if col_data(k)]
    column_groups = [{'group_title': '', 'columns': all_columns}]

    return render(request, 'proposals/kanban.html', {
        'column_groups': column_groups,
    })


@login_required
@require_http_methods(['GET'])
def kanban_card_detail_view(request, tkp_id):
    """Данные карточки для модального окна канбана: ТКП + договор (если есть) + доп. поля."""
    try:
        tkp = TKPRecord.objects.get(pk=tkp_id)
    except TKPRecord.DoesNotExist:
        raise Http404()
    contract = ContractRecord.objects.filter(tkp_id=tkp_id).select_related('counterparty').first()
    service_display = 'Комплексное' if tkp.service == 'Комплексное ТКП' else tkp.service
    custom_fields = list(
        KanbanCardField.objects.filter(user=request.user, tkp_id=tkp_id).order_by('order', 'pk')
    )
    ctx = {
        'tkp': tkp,
        'contract': contract,
        'tkp_date_display': tkp.date.strftime('%d.%m.%Y'),
        'service_display': service_display,
        'sum_display': _format_price(tkp.sum_total or 0),
        'custom_fields': custom_fields,
    }
    if contract:
        ctx['contract_date_display'] = contract.date.strftime('%d.%m.%Y')
        ctx['contract_sum_display'] = _format_price(contract.sum_total or 0)
    return render(request, 'proposals/kanban_card_content.html', ctx)


@login_required
@require_http_methods(['POST'])
def kanban_column_title_view(request):
    """Сохранение переименования колонки канбана (стандартной или пользовательской)."""
    column_key = (request.POST.get('column_key') or '').strip()
    title = (request.POST.get('title') or '').strip()
    if not title:
        return JsonResponse({'success': False, 'error': 'empty_title'}, status=400)
    if column_key.startswith('custom_'):
        try:
            col_id = int(column_key.replace('custom_', ''))
            col = KanbanColumnCustom.objects.get(pk=col_id, user=request.user)
            col.title = title
            col.save(update_fields=['title'])
            return JsonResponse({'success': True})
        except (ValueError, KanbanColumnCustom.DoesNotExist):
            return JsonResponse({'success': False, 'error': 'not_found'}, status=404)
    allowed = (KANBAN_COL_DRAFT, KANBAN_COL_FINAL, KANBAN_COL_CONTRACT_DRAFT, KANBAN_COL_CONTRACT_FINAL)
    if column_key not in allowed:
        return JsonResponse({'success': False, 'error': 'invalid_column'}, status=400)
    obj, _ = KanbanColumnTitleOverride.objects.update_or_create(
        user=request.user,
        column_key=column_key,
        defaults={'title': title},
    )
    return JsonResponse({'success': True})


@login_required
@require_http_methods(['POST'])
def kanban_column_reorder_view(request):
    """Сохранение порядка колонок на канбане (список ключей колонок)."""
    try:
        body = request.body.decode('utf-8') if isinstance(request.body, bytes) else (request.body or '[]')
        order = json.loads(body) if body else []
    except (ValueError, TypeError, UnicodeDecodeError):
        order = []
    if not isinstance(order, list):
        return JsonResponse({'success': False, 'error': 'invalid'}, status=400)
    order = [str(k) for k in order if k]
    KanbanBoardOrder.objects.update_or_create(
        user=request.user,
        defaults={'order': order},
    )
    return JsonResponse({'success': True})


@login_required
@require_http_methods(['POST'])
def kanban_column_create_view(request):
    """Создание пользовательской колонки канбана."""
    title = (request.POST.get('title') or '').strip()
    if not title:
        return JsonResponse({'success': False, 'error': 'empty_title'}, status=400)
    max_order = KanbanColumnCustom.objects.filter(user=request.user).aggregate(
        m=Max('order')
    )['m'] or 0
    col = KanbanColumnCustom.objects.create(user=request.user, title=title, order=max_order + 1)
    return JsonResponse({'success': True, 'column_id': col.pk, 'column_key': f'custom_{col.pk}'})


@login_required
@require_http_methods(['POST'])
def kanban_card_move_view(request):
    """Перемещение карточки в колонку (сохранение позиции на канбане)."""
    tkp_id = request.POST.get('tkp_id')
    column_key = (request.POST.get('column_key') or '').strip()
    if not tkp_id or not column_key:
        return JsonResponse({'success': False, 'error': 'missing_params'}, status=400)
    try:
        tkp_id = int(tkp_id)
    except (ValueError, TypeError):
        return JsonResponse({'success': False, 'error': 'invalid_tkp_id'}, status=400)
    if not TKPRecord.objects.filter(pk=tkp_id).exists():
        return JsonResponse({'success': False, 'error': 'not_found'}, status=404)
    if column_key.startswith('custom_'):
        try:
            col_id = int(column_key.replace('custom_', ''))
            if not KanbanColumnCustom.objects.filter(pk=col_id, user=request.user).exists():
                return JsonResponse({'success': False, 'error': 'invalid_column'}, status=400)
        except ValueError:
            return JsonResponse({'success': False, 'error': 'invalid_column'}, status=400)
    else:
        allowed = (KANBAN_COL_DRAFT, KANBAN_COL_FINAL, KANBAN_COL_CONTRACT_DRAFT, KANBAN_COL_CONTRACT_FINAL)
        if column_key not in allowed:
            return JsonResponse({'success': False, 'error': 'invalid_column'}, status=400)
    KanbanCardPosition.objects.update_or_create(
        user=request.user,
        tkp_id=tkp_id,
        defaults={'column_key': column_key},
    )
    return JsonResponse({'success': True})


@login_required
@require_http_methods(['POST'])
def kanban_card_field_save_view(request, tkp_id):
    """Добавление или обновление доп. поля карточки канбана."""
    try:
        tkp = TKPRecord.objects.get(pk=tkp_id)
    except TKPRecord.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'not_found'}, status=404)
    name = (request.POST.get('name') or '').strip()
    value_type = (request.POST.get('value_type') or 'text').strip().lower()
    value = (request.POST.get('value') or '').strip()
    if not name:
        return JsonResponse({'success': False, 'error': 'empty_name'}, status=400)
    if value_type not in (KanbanCardField.VALUE_TEXT, KanbanCardField.VALUE_NUMBER):
        value_type = KanbanCardField.VALUE_TEXT
    try:
        field = KanbanCardField.objects.get(user=request.user, tkp_id=tkp_id, name=name)
        field.value_type = value_type
        field.value = value
        field.save(update_fields=['value_type', 'value'])
    except KanbanCardField.DoesNotExist:
        max_order = KanbanCardField.objects.filter(user=request.user, tkp_id=tkp_id).aggregate(m=Max('order'))['m'] or 0
        field = KanbanCardField.objects.create(
            user=request.user,
            tkp_id=tkp_id,
            name=name,
            value_type=value_type,
            value=value,
            order=max_order + 1,
        )
    return JsonResponse({'success': True, 'field_id': field.pk})


@login_required
@require_http_methods(['POST'])
def kanban_save_notes_view(request, tkp_id):
    """Сохранение заметок по сделке (карточка канбана)."""
    try:
        tkp = TKPRecord.objects.get(pk=tkp_id)
    except TKPRecord.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'not_found'}, status=404)
    notes = (request.POST.get('notes') or '').strip()
    tkp.notes = notes
    tkp.save(update_fields=['notes'])
    return JsonResponse({'success': True})


@login_required
@require_http_methods(['GET', 'POST'])
def contract_table_view(request):
    """Реестр договоров: фильтры (клиент, услуга, статус), сортировка по всем колонкам, колонка Статус договора; удаление записей."""
    if request.method == 'POST':
        q = request.GET.urlencode()
        url = reverse('proposals:contract_table') + ('?' + q if q else '')
        bulk_ids = request.POST.getlist('ids')
        bulk_action = (request.POST.get('bulk_action') or '').strip()

        if bulk_action == 'delete' and bulk_ids:
            deleted = 0
            for pk in bulk_ids:
                try:
                    rec = ContractRecord.objects.get(pk=pk)
                    base_name = rec.docx_file or rec.pdf_file
                    rec.delete()
                    if base_name:
                        _delete_tkp_files(base_name)
                    deleted += 1
                except (ContractRecord.DoesNotExist, ValueError, TypeError):
                    pass
            if deleted:
                messages.success(request, f'Удалено записей: {deleted}.')
            return redirect(url)

        if bulk_action == 'copy' and bulk_ids:
            copied = 0
            for pk in bulk_ids:
                try:
                    rec = ContractRecord.objects.get(pk=pk)
                    seq = _get_next_contract_seq_for_date(rec.date)
                    new_number = f'{rec.date:%d%m%Y}_{seq}'
                    ContractRecord.objects.create(
                        date=rec.date,
                        number=new_number,
                        status=ContractRecord.STATUS_DRAFT,
                        tkp=rec.tkp,
                        counterparty=rec.counterparty,
                        client=rec.client or '',
                        service=rec.service or '',
                        sum_total=rec.sum_total or Decimal(0),
                        docx_file='',
                        pdf_file='',
                        contract_snapshot=rec.contract_snapshot,
                        created_by=request.user,
                    )
                    copied += 1
                except (ContractRecord.DoesNotExist, ValueError, TypeError):
                    pass
            if copied:
                messages.success(request, f'Скопировано записей: {copied}.')
            return redirect(url)

        delete_id = request.POST.get('delete_id', '').strip()
        if delete_id:
            try:
                rec = ContractRecord.objects.get(pk=delete_id)
                base_name = rec.docx_file or rec.pdf_file
                rec.delete()
                if base_name:
                    _delete_tkp_files(base_name)
                messages.success(request, 'Запись договора удалена.')
            except ContractRecord.DoesNotExist:
                messages.error(request, 'Запись не найдена.')
            return redirect(url)
    records = ContractRecord.objects.select_related('tkp', 'counterparty', 'created_by')
    client = request.GET.get('client', '').strip()
    service = request.GET.get('service', '').strip()
    status_filter = request.GET.get('status', '').strip().lower()
    if client:
        records = records.filter(client__icontains=client)
    if service:
        records = records.filter(service=service)
    if status_filter == 'draft':
        records = records.filter(status=ContractRecord.STATUS_DRAFT)
    elif status_filter == 'final':
        records = records.filter(status=ContractRecord.STATUS_FINAL)
    sort_col = request.GET.get('sort', '').strip().lower()
    order = request.GET.get('order', 'desc').strip().lower()
    if order not in ('asc', 'desc'):
        order = 'desc'
    order_prefix = '' if order == 'asc' else '-'
    allowed_sort = {'date': 'date', 'number': 'number', 'client': 'client', 'service': 'service', 'sum_total': 'sum_total', 'status': 'status'}
    if sort_col in allowed_sort:
        field = allowed_sort[sort_col]
        if field == 'status':
            records = records.order_by(f'{order_prefix}status', '-date', '-created_at')
        else:
            records = records.order_by(f'{order_prefix}{field}', '-created_at')
    else:
        records = records.order_by('-date', '-created_at')
    services_list = list(ContractRecord.objects.values_list('service', flat=True).distinct().filter(service__isnull=False).exclude(service='').order_by('service'))
    context = {
        'records': records,
        'filters': {'client': client, 'service': service, 'status': status_filter},
        'services_list': services_list,
        'sort': sort_col or 'date',
        'order': order,
    }
    return render(request, 'proposals/contract_table.html', context)


@login_required
@require_http_methods(['GET'])
def contract_card_view(request, contract_id):
    """Карточка договора (фрагмент для модального окна): реквизиты только для просмотра."""
    try:
        contract = ContractRecord.objects.select_related('counterparty', 'tkp').get(pk=contract_id)
    except ContractRecord.DoesNotExist:
        raise Http404()
    snapshot = contract.contract_snapshot or {}
    _cp = contract.counterparty
    ctx = {
        'contract': contract,
        'date_display': contract.date.strftime('%d.%m.%Y'),
        'customer_name': snapshot.get('customer_name') or (_cp.name if _cp else '') or contract.client,
        'customer_represented_by': snapshot.get('customer_represented_by') or (_cp.director and _director_genitive(_cp.director) or '') if _cp else '',
        'customer_represented_by_nominative': snapshot.get('customer_represented_by_nominative') or (_cp.director if _cp else ''),
        'price': snapshot.get('price') or _format_price(contract.sum_total),
        'payment_terms': snapshot.get('payment_terms') or '',
        'name': snapshot.get('name') or (_cp.name if _cp else ''),
        'address': snapshot.get('address') or (_cp.address if _cp else ''),
        'inn': snapshot.get('inn') or (_cp.inn if _cp else ''),
        'kpp': snapshot.get('kpp') or (_cp.kpp if _cp else ''),
        'ogrn': snapshot.get('ogrn') or (_cp.ogrn if _cp else ''),
        'account': snapshot.get('account') or (_cp.account if _cp else ''),
        'bank': snapshot.get('bank') or (_cp.bank if _cp else ''),
        'bik': snapshot.get('bik') or (_cp.bik if _cp else ''),
        'kor_account': snapshot.get('kor_account') or (_cp.kor_account if _cp else ''),
        'email': snapshot.get('email') or (_cp.email if _cp else ''),
        'room': snapshot.get('room') or '',
        's': snapshot.get('s') or '',
    }
    return render(request, 'proposals/contract_card_content.html', ctx)


@login_required
@require_http_methods(['GET'])
def contract_download_file_view(request, contract_id, file_type):
    """Скачивание файла договора (docx или pdf) по ID записи."""
    if file_type not in ('docx', 'pdf'):
        raise Http404()
    try:
        contract = ContractRecord.objects.get(pk=contract_id)
    except ContractRecord.DoesNotExist:
        raise Http404()
    base_name = contract.docx_file if file_type == 'docx' else contract.pdf_file
    if not base_name:
        raise Http404()
    out_dir = Path(getattr(settings, 'TKP_OUTPUT_DIR', settings.BASE_DIR / 'TKP_output'))
    ext = 'pdf' if file_type == 'pdf' else 'docx'
    path = out_dir / f'{base_name}.{ext}'
    if not path.exists():
        raise Http404()
    return FileResponse(
        open(path, 'rb'),
        as_attachment=True,
        filename=path.name,
    )


@login_required
@require_http_methods(['GET', 'POST'])
def counterparties_view(request):
    """Список контрагентов (реквизиты из карточки). Фильтры по наименованию и ИНН; удаление."""
    if request.method == 'POST':
        delete_id = request.POST.get('delete_id', '').strip()
        if delete_id:
            try:
                Counterparty.objects.filter(pk=delete_id).delete()
                messages.success(request, 'Контрагент удалён.')
            except Exception:
                pass
            q = request.GET.urlencode()
            url = reverse('proposals:counterparties') + ('?' + q if q else '')
            return redirect(url)
    records = Counterparty.objects.all()
    name_filter = request.GET.get('name', '').strip()
    inn_filter = request.GET.get('inn', '').strip()
    if name_filter:
        records = records.filter(name__icontains=name_filter)
    if inn_filter:
        records = records.filter(inn__icontains=inn_filter)
    return_url = request.GET.get('return_url', '').strip()
    context = {
        'records': records,
        'filters': {'name': name_filter, 'inn': inn_filter},
        'return_url': return_url,
    }
    return render(request, 'proposals/counterparties.html', context)


@login_required
@require_http_methods(['GET'])
def counterparty_search_view(request):
    """Поиск контрагентов по наименованию и ИНН (JSON для автоподстановки в форме договора)."""
    q = (request.GET.get('q') or '').strip()
    if not q or len(q) < 2:
        return JsonResponse({'results': []})
    from django.db.models import Q
    qs = Counterparty.objects.filter(
        Q(name__icontains=q) | Q(inn__icontains=q)
    ).order_by('name')[:20]
    results = [{'id': c.pk, 'name': c.name or '', 'inn': c.inn or ''} for c in qs]
    return JsonResponse({'results': results})


@login_required
@require_http_methods(['GET'])
def counterparty_json_view(request, pk):
    """Реквизиты контрагента по ID (JSON для подстановки в форму договора)."""
    try:
        cp = Counterparty.objects.get(pk=pk)
    except Counterparty.DoesNotExist:
        return JsonResponse({}, status=404)
    return JsonResponse({
        'name': cp.name or '',
        'inn': cp.inn or '',
        'kpp': cp.kpp or '',
        'address': cp.address or '',
        'director': cp.director or '',
        'director_genitive': _director_genitive(cp.director or ''),
        'ogrn': cp.ogrn or '',
        'account': cp.account or '',
        'bank': cp.bank or '',
        'bik': cp.bik or '',
        'kor_account': cp.kor_account or '',
        'email': cp.email or '',
    })


@login_required
@require_http_methods(['GET', 'POST'])
def tariffs_view(request):
    """Редактирование тарифов по регионам: список пар Услуга–Регион и форма добавления."""
    if request.method == 'POST':
        delete_id = request.POST.get('delete_id')
        if delete_id:
            try:
                RegionServicePrice.objects.filter(pk=delete_id).delete()
                messages.success(request, 'Тариф удалён.')
            except Exception:
                pass
            return redirect('proposals:tariffs')
        form = TariffForm(request.POST)
        if form.is_valid():
            data = form.cleaned_data
            service = data['service']
            new_name = (data.get('new_region_name') or '').strip()
            if new_name:
                region, _ = Region.objects.get_or_create(name=new_name, defaults={})
            else:
                region = data['region']
            RegionServicePrice.objects.update_or_create(
                region=region,
                service=service,
                defaults={'unit_price': data['unit_price']},
            )
            messages.success(request, f'Тариф сохранён: {service.name} — {region.name}.')
            return redirect('proposals:tariffs')
    else:
        form = TariffForm()
    tariffs_list = RegionServicePrice.objects.select_related('region', 'service').order_by('region__name', 'service__name')
    context = {'form': form, 'tariffs_list': tariffs_list}
    return render(request, 'proposals/tariffs.html', context)


@login_required
@require_http_methods(['GET', 'POST'])
def service_descriptions_view(request):
    """Редактирование описаний услуг. Описание подставляется в поле «Комментарий» в Комплексном ТКП."""
    services_qs = Service.objects.order_by('order', 'name')
    if request.method == 'POST':
        for service in services_qs:
            key = f'desc_{service.id}'
            new_desc = (request.POST.get(key) or '').strip()
            if service.description != new_desc:
                service.description = new_desc
                service.save(update_fields=['description'])
        # Сохраняем в файл, чтобы при деплое не вводить комментарии заново
        try:
            updated = Service.objects.order_by('order', 'name').values('name', 'description')
            comments_by_name = {s['name']: (s['description'] or '') for s in updated}
            _save_complex_service_comments_file(comments_by_name)
        except Exception:
            pass
        messages.success(request, 'Описания услуг сохранены.')
        return redirect('proposals:service_descriptions')
    file_comments = _load_complex_service_comments_file()
    services = []
    for s in services_qs:
        display_name = COMPLEX_SERVICE_DISPLAY_NAMES.get(s.name, s.name)
        desc = (s.description or '').strip()
        if not desc:
            desc = file_comments.get(s.name, '')
        services.append({
            'id': s.id,
            'name': s.name,
            'display_name': display_name,
            'description': desc,
        })
    context = {'services': services}
    return render(request, 'proposals/service_descriptions.html', context)


@login_required
@require_http_methods(['GET', 'POST'])
def requisites_add_view(request):
    """Загрузка карточки контрагента, извлечение реквизитов и формирование пользовательской карточки."""
    form = RequisitesParseForm()
    card_data = None

    if request.method == 'POST':
        action = (request.POST.get('action') or '').strip()
        form = RequisitesParseForm(request.POST, request.FILES)

        if action == 'parse':
            source_file = request.FILES.get('source_file')
            if not source_file:
                form.add_error('source_file', 'Выберите файл для извлечения реквизитов.')
            elif form.is_valid():
                try:
                    parsed = parse_requisites_file(source_file.name, source_file.read())
                except ValueError as exc:
                    messages.error(request, str(exc))
                except Exception:
                    messages.error(
                        request,
                        'Не удалось обработать файл. Проверьте, что это корректный DOCX/PDF с текстом.',
                    )
                else:
                    merged = {}
                    for field in FIELD_ORDER:
                        parsed_value = (parsed.get(field) or '').strip()
                        manual_value = (form.cleaned_data.get(field) or '').strip()
                        merged[field] = parsed_value or manual_value
                    form = RequisitesParseForm(initial=merged)
                    messages.success(request, 'Реквизиты извлечены. Проверьте и при необходимости отредактируйте.')

        elif action == 'build':
            if form.is_valid():
                card_data = {field: (form.cleaned_data.get(field) or '').strip() for field in FIELD_ORDER}
                if not any(card_data.values()):
                    form.add_error(None, 'Заполните хотя бы одно поле реквизитов, чтобы сформировать карточку.')
                    card_data = None
                else:
                    inn = card_data.get('inn', '').strip()
                    if inn and Counterparty.objects.filter(inn=inn).exists():
                        messages.warning(request, 'Контрагент уже заведён (такой ИНН есть в таблице Контрагенты).')
                    else:
                        new_cp = Counterparty.objects.create(**card_data)
                        tkp_id = request.GET.get('tkp_id')
                        if request.GET.get('from') == 'contract' and tkp_id:
                            messages.success(request, 'Карточка создана. Перейдите к вводу данных договора.')
                            url = reverse('proposals:contract_form', args=[tkp_id]) + '?counterparty_id=' + str(new_cp.pk)
                            return redirect(url)
                        messages.success(request, 'Карточка создана. Контрагент добавлен в таблицу Контрагенты.')
            else:
                messages.error(request, 'Проверьте корректность заполнения полей.')

    tkp_id = request.GET.get('tkp_id')
    from_contract = request.GET.get('from') == 'contract'
    context = {
        'form': form,
        'card_data': card_data,
        'tkp_id': tkp_id,
        'from_contract': from_contract,
    }
    return render(request, 'proposals/requisites_form.html', context)


def _generate_doc_number(client, date_obj, seq):
    """Генерация номера документа: client_DDMMYYYY_N."""
    client_safe = _sanitize_filename(client or '')
    date_str = date_obj.strftime('%d%m%Y')
    return f'{client_safe}_{date_str}_{seq}'


def _director_genitive(director):
    """ФИО директора в родительном падеже (для «Заказчик в лице … действующего»)."""
    if not director or not (director or '').strip():
        return ''
    try:
        import pymorphy2
        morph = pymorphy2.MorphAnalyzer()
        parts = (director or '').strip().split()
        result = []
        for word in parts:
            parsed = morph.parse(word)
            if not parsed:
                result.append(word)
                continue
            p = parsed[0]
            inflected = p.inflect({'gent'}) if p.tag.case else None
            result.append(inflected.word if inflected else word)
        return ' '.join(result)
    except Exception:
        return (director or '').strip()


def _get_next_draft_seq_for_date(date_obj):
    """Порядковый номер черновика за указанную дату."""
    return TKPRecord.objects.filter(date=date_obj, status=TKPRecord.STATUS_DRAFT).count() + 1


def _generate_draft_number(date_obj, seq):
    """Номер черновика: Черновик_DDMMYYYY_N."""
    return f'Черновик_{date_obj:%d%m%Y}_{seq}'


def _save_tkp_record(data, status=None, user=None):
    """Сохранение записи о сформированном ТКП (status по умолчанию — итоговый)."""
    from datetime import datetime
    date_obj = datetime.strptime(data['date'], '%Y-%m-%d').date()
    if status == TKPRecord.STATUS_DRAFT:
        seq = _get_next_draft_seq_for_date(date_obj)
        number = _generate_draft_number(date_obj, seq)
    else:
        seq = _get_next_seq_for_date(date_obj)
        number = _generate_doc_number(data.get('client') or '', date_obj, seq)
    TKPRecord.objects.create(
        date=date_obj,
        number=number,
        client=data.get('client') or '',
        service=data['service_name'],
        sum_total=Decimal(data.get('price') or 0),
        room=data.get('room') or '',
        s=str(data.get('s') or ''),
        text=data.get('text') or '',
        status=status or TKPRecord.STATUS_FINAL,
        created_by=user,
    )


def _generate_and_save_files(data):
    """Генерация docx, конвертация в PDF, сохранение обоих в TKP_output. Возвращает base_name файлов."""
    try:
        service = Service.objects.get(pk=data['service_id'])
    except Service.DoesNotExist:
        return None

    templates_dir = getattr(settings, 'TEMPLATES_DOCX_DIR', Path(settings.BASE_DIR) / 'templates_docx')
    template_path = templates_dir / service.template_file

    if not template_path.exists():
        return None

    from datetime import datetime
    date_obj = datetime.strptime(data['date'], '%Y-%m-%d').date()
    date_display = date_obj.strftime('%d.%m.%Y')
    seq = _get_next_seq_for_date(date_obj)
    number = f'{date_obj:%d%m%Y}_{seq}'

    price_val = Decimal(data['price'])

    context = {
        'city': data.get('city', ''),
        'price': _format_price(price_val),
        'date': date_display,
        'client': data['client'] or '',
        'room': data.get('room') or '',
        'srok': data.get('srok') or '',
        'text': data['text'] or '',
        's': data.get('s') or '',
        'number': number,
    }

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir = Path(tmpdir)
        docx_path = tmpdir / 'tkp.docx'
        pdf_path = tmpdir / 'tkp.pdf'

        doc = DocxTemplate(str(template_path))
        doc.render(context)
        doc.save(str(docx_path))

        _convert_docx_to_pdf(docx_path, pdf_path.parent)

        client_safe = _sanitize_filename(data.get('client') or '')
        base_name = f'{client_safe}_{date_obj:%d%m%Y}_{seq}'
        out_dir = Path(getattr(settings, 'TKP_OUTPUT_DIR', settings.BASE_DIR / 'TKP_output'))
        out_dir.mkdir(parents=True, exist_ok=True)
        shutil.copy2(pdf_path, out_dir / f'{base_name}.pdf')
        shutil.copy2(docx_path, out_dir / f'{base_name}.docx')
        return base_name
