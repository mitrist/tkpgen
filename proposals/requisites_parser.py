import re
from io import BytesIO

from docx import Document


FIELD_ORDER = (
    'name',
    'inn',
    'address',
    'director',
    'ogrn',
    'account',
    'bank',
    'bik',
    'phone',
)


def parse_requisites_file(file_name, file_bytes):
    """Возвращает словарь реквизитов из DOCX/PDF."""
    ext = (file_name.rsplit('.', 1)[-1] if '.' in file_name else '').lower()
    if ext == 'docx':
        text = _extract_text_from_docx(file_bytes)
    elif ext == 'pdf':
        text = _extract_text_from_pdf(file_bytes)
    else:
        raise ValueError('Поддерживаются только файлы .docx и .pdf')
    return _extract_requisites(text)


def _extract_text_from_docx(file_bytes):
    doc = Document(BytesIO(file_bytes))
    lines = []
    for paragraph in doc.paragraphs:
        _append_line(lines, paragraph.text)
    for table in doc.tables:
        _collect_table_lines(table, lines)
    return '\n'.join(lines)


def _collect_table_lines(table, lines):
    for row in table.rows:
        for cell in row.cells:
            cell_text = ' '.join(p.text.strip() for p in cell.paragraphs if p.text and p.text.strip())
            _append_line(lines, cell_text)
            for nested in cell.tables:
                _collect_table_lines(nested, lines)


def _extract_text_from_pdf(file_bytes):
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError(
            'Для обработки PDF установите пакет pypdf: pip install pypdf'
        ) from exc
    reader = PdfReader(BytesIO(file_bytes))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or '')
    return '\n'.join(parts)


def _append_line(lines, value):
    value = _clean_value(value)
    if value:
        lines.append(value)


def _normalize_text(text):
    if not text:
        return ''
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = text.replace('\xa0', ' ')
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{2,}', '\n', text)
    return text.strip()


def _clean_value(value):
    if not value:
        return ''
    value = value.replace('\xa0', ' ')
    value = re.sub(r'[ \t]+', ' ', value).strip(' \t\n-:;,.')
    return value


def _extract_requisites(raw_text):
    text = _normalize_text(raw_text)
    lines = [_clean_value(line) for line in text.split('\n') if _clean_value(line)]
    result = {k: '' for k in FIELD_ORDER}

    result['inn'] = _search_first(
        text,
        (
            r'\bИНН\s*[:\-]?\s*(\d{10,12})\b',
            r'\bИНН\/КПП\s*[:\-]?\s*(\d{10,12})\b',
        ),
    )
    result['ogrn'] = _search_first(text, (r'\bОГРН\s*[:\-]?\s*(\d{13,15})\b',))
    result['bik'] = _search_first(text, (r'\bБИК\s*[:\-]?\s*(\d{9})\b',))
    result['account'] = _search_first(
        text,
        (
            r'(?:р\/сч(?:е[тт])?|р\/счет|расч[её]тн(?:ый|ого)\s+сч[её]т)\s*[:\-]?\s*([0-9 ]{20,30})',
            r'\b([0-9]{20})\b',
        ),
        postprocess_digits=True,
    )
    result['phone'] = _search_first(
        text,
        (
            r'\b(?:тел(?:ефон)?\.?)\s*[:\-]?\s*([\+\d\(\)\s\-]{7,30})',
            r'((?:\+7|8)\s*\(?\d{3,5}\)?[\s\-]*\d[\d\s\-]{5,20})',
        ),
    )
    result['director'] = _search_first(
        text,
        (
            r'(?:генеральн\w*\s+директор|директор)\s*[:\-–]\s*([^\n,]+)',
            r'(?:генеральн\w*\s+директор|директор)\s+([А-ЯЁA-Z][^\n,]{3,120})',
        ),
    )

    result['name'] = _find_value_by_labels(
        lines,
        ('полное наименование', 'наименование организации', 'наименование контрагента', 'наименование'),
    )
    if not result['name']:
        result['name'] = _search_first(
            text,
            (
                r'^\s*((?:ООО|АО|ПАО|ОАО|ЗАО|ИП|МБОУ|МОУ|ГБУ|ФГБУ)[^\n]{3,300})',
                r'^\s*["«]([^"\n»]{5,300})["»]',
            ),
            flags=re.IGNORECASE | re.MULTILINE,
        )
    if not result['name']:
        for line in lines[:8]:
            low = line.lower()
            if 'реквизит' in low or 'карта предприятия' in low:
                continue
            if len(line) > 8:
                result['name'] = line
                break

    result['address'] = _find_value_by_labels(lines, ('адрес', 'юридический адрес', 'местонахождение'))
    result['bank'] = _find_value_by_labels(
        lines,
        ('наименование банка', 'банк получателя', 'в банке'),
        min_len=4,
    )
    if not result['bank']:
        for line in lines:
            low = line.lower()
            if 'банк' in low and 'бик' not in low:
                candidate = re.sub(
                    r'^(?:наименование\s+банка|банк(?:\s+получателя)?|в\s+банке)\s*[:\-]?\s*',
                    '',
                    line,
                    flags=re.IGNORECASE,
                )
                candidate = _clean_value(candidate)
                if len(candidate) >= 4:
                    result['bank'] = candidate
                    break

    for key in FIELD_ORDER:
        result[key] = _clean_value(result[key])
    return result


def _find_value_by_labels(lines, labels, min_len=1):
    labels_l = tuple(l.lower() for l in labels)
    for idx, line in enumerate(lines):
        low = line.lower()
        for label in labels_l:
            if label not in low:
                continue
            value = _extract_inline_after_label(line, label)
            if value and len(value) >= min_len:
                return value
            if idx + 1 < len(lines):
                candidate = lines[idx + 1]
                if not _looks_like_label(candidate) and len(candidate) >= min_len:
                    return candidate
    return ''


def _extract_inline_after_label(line, label):
    low = line.lower()
    pos = low.find(label)
    if pos < 0:
        return ''
    tail = line[pos + len(label):].strip()
    tail = tail.lstrip(':-– ')
    return _clean_value(tail)


def _looks_like_label(line):
    low = line.lower()
    if ':' in line and len(line) < 80:
        return True
    markers = (
        'инн',
        'кпп',
        'огрн',
        'бик',
        'тел',
        'адрес',
        'директор',
        'банк',
        'счет',
        'р/сч',
    )
    return any(marker in low for marker in markers)


def _search_first(text, patterns, flags=re.IGNORECASE, postprocess_digits=False):
    for pattern in patterns:
        match = re.search(pattern, text, flags)
        if not match:
            continue
        value = match.group(1)
        value = _clean_value(value)
        if postprocess_digits:
            digits = re.sub(r'\D', '', value)
            if len(digits) >= 20:
                return digits[:20]
        if value:
            return value
    return ''
